# -*- coding: utf-8 -*-
"""v25: v23-аас эхлээд

1. Caption font 10pt, bold=False (body reference-г хөндөхгүй)
2. 5 зургийг шинэ, эх сурвалж хадгалагдаагүй PNG-ээр солих
3. Текст дэх "Эх сурвалж: Оюутны тооцоолол" мөрүүдийг ҮЛДЭЭНЭ
"""
import sys
import io
import os
import re
import shutil
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

src = 'outputs/paper/bagiin_ner_v23.docx'
dst = 'outputs/paper/bagiin_ner_v25.docx'
shutil.copy(src, dst)
d = Document(dst)

FIG_DIR = 'outputs/figures'
FIG_WIDTH = Inches(6.0)

# ==================== 1) Caption форматыг 10pt, non-bold болгох ==============
cap_count = 0
for p in d.paragraphs:
    t = p.text.strip()
    if not t:
        continue
    # Body caption нь "Хүснэгт N " эсвэл "Зураг N " (period ба tab байхгүй)
    is_cap = False
    for prefix in ('Хүснэгт ', 'Зураг '):
        if t.startswith(prefix):
            head = t[:12]
            if '.' not in head and '\t' not in head:
                # Body ref-ийг хасах ("Хүснэгт 4-өөс...")
                if not re.match(r'^(Хүснэгт|Зураг)\s+\d+-', t):
                    is_cap = True
            break
    if not is_cap:
        continue
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.bold = False
    cap_count += 1

print(f"Зассан caption: {cap_count}")

# ==================== 2) 5 зургийн PNG-ийг шинэ хувилбараар солих =========
# Одоогийн image paragraph-уудыг олох
image_paras = []
for i, p in enumerate(d.paragraphs):
    if p._element.findall('.//' + qn('w:drawing')):
        image_paras.append(i)

print(f"Зурагтай parа: {image_paras}")


def insert_picture_before(target_el, image_path, width):
    new_p = d.add_paragraph()
    run = new_p.add_run()
    run.add_picture(image_path, width=width)
    new_el = new_p._element
    new_el.getparent().remove(new_el)
    target_el.addprevious(new_el)


fig_files = [
    'f1_education_distribution.png',
    'f2_wage_education_scatter.png',
    'f3_threshold_profile.png',
    'f4_regime_slopes.png',
    'f5_ols_vs_iv.png',
]

# Хуучныг устгаж, яг тэнд нь шинийг оруулах
# Буцаасан дараалалтай устгаж, өмнө нь шинийг оруулах
for idx in reversed(range(len(image_paras))):
    para_idx = image_paras[idx]
    fig_name = fig_files[idx]
    img_path = os.path.join(FIG_DIR, fig_name)

    old_p = d.paragraphs[para_idx]
    old_el = old_p._element
    # Хуучны өмнө шинэ paragraph оруулах
    insert_picture_before(old_el, img_path, FIG_WIDTH)
    # Хуучныг устгах
    old_el.getparent().remove(old_el)
    print(f"  Replaced [{para_idx}] -> {fig_name}")

d.save(dst)
print(f"\nSaved: {dst}")

# ==================== Баталгаажуулалт ================================
d2 = Document(dst)
# Эх сурвалж үлдсэн эсэх
remain_src = sum(1 for p in d2.paragraphs
                 if p.text.strip().startswith('Эх сурвалж'))
print(f"Үлдсэн 'Эх сурвалж' мөр: {remain_src} (13 байх ёстой)")

# Зургийн тоо
total_images = 0
for p in d2.paragraphs:
    if p._element.findall('.//' + qn('w:drawing')):
        total_images += 1
print(f"Зурагийн тоо: {total_images} (5 байх ёстой)")
