# -*- coding: utf-8 -*-
"""v14: хиймэл → псевдо буцаах + α_{ca} LaTeX тэмдэглэгээ."""
import sys
import io
import shutil
from docx import Document

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

src = 'outputs/paper/bagiin_ner_v13.docx'
dst = 'outputs/paper/bagiin_ner_v14.docx'
shutil.copy(src, dst)
d = Document(dst)

replacements = [
    ('хиймэл панелийн', 'псевдо-панелийн'),
    ('хиймэл панел',    'псевдо-панел'),
    ('Хиймэл панел',    'Псевдо-панел'),
]

changed = 0
for p in d.paragraphs:
    for run in p.runs:
        orig = run.text
        new = orig
        for old, repl in replacements:
            new = new.replace(old, repl)
        if new != orig:
            run.text = new
            changed += 1


def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


new_99 = (
    "Энд α_{ca} нь c когорт болон a аймгийн нүдэнд харгалзах тогтмол "
    "хүчин зүйл (ажиглагдашгүй орон нутаг, үеийн шинж), τ_t нь t "
    "давалгааны тогтмол хүчин зүйл, ε̄_{ca,t} нь нүдний санамсаргүй "
    "алдаа юм. Тогтмол нөлөөний загварт α_{ca}-г тайлбарлах хувьсагчтай "
    "хамааралтай тогтмол параметр гэж үздэг бол санамсаргүй нөлөөний "
    "загвар нь α_{ca} ~ N(0, σ_α^2) хэлбэрийн санамсаргүй хэмжигдэхүүн "
    "хэмээн үзнэ. Эдгээр загваруудын аль нь илүү тохиромжтой болохыг "
    "Хаусманы шалгуураар тогтоов:"
)
set_text(d.paragraphs[99], new_99)

d.save(dst)
print(f"Saved: {dst}, {changed} run-ийг засав")

d2 = Document(dst)
for i in [91, 97, 99, 193]:
    print("\n[{}]: {}".format(i, d2.paragraphs[i].text[:400]))

for i, p in enumerate(d2.paragraphs):
    if 'хиймэл' in p.text.lower():
        print("REMAINS [{}]: {}".format(i, p.text[:150]))
