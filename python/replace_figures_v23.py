# -*- coding: utf-8 -*-
"""v23: 5 зургийг шинэчилж солих, F5 нэмэх, товчилсон үгсийн хүснэгтийг шинэчлэх."""
import sys
import io
import shutil
import os
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

src = 'outputs/paper/bagiin_ner_v22.docx'
dst = 'outputs/paper/bagiin_ner_v23.docx'
shutil.copy(src, dst)
d = Document(dst)

FIG_DIR = 'outputs/figures'
FIG_WIDTH = Inches(6.0)

# ==================== 1) Зургийн байрлалуудыг олох ==================
# Одоогийн state: [152] F1, [156] F2, [177] F3, [184] F4 — inline images
# F5 байхгүй, Зураг 5 caption [194]-д байгаа

paras = list(d.paragraphs)
image_locations = []  # (para_index, fig_name)
for i, p in enumerate(paras):
    imgs = p._element.findall('.//' + qn('w:drawing'))
    if imgs:
        image_locations.append(i)

print(f"Одоогийн зураг бүхий parа: {image_locations}")

fig_files = {
    0: 'f1_education_distribution.png',
    1: 'f2_wage_education_scatter.png',
    2: 'f3_threshold_profile.png',
    3: 'f4_regime_slopes.png',
}

# ==================== 2) Байгаа 4 зургийг шинэчлэх ==================
# Арга: хэрэв зураг парa-ийн дараа оруулна, дараа нь хуучинг устгах

def insert_picture_after(target_el, image_path, width):
    """add_picture with reparent to target_el's next."""
    new_p = d.add_paragraph()
    run = new_p.add_run()
    run.add_picture(image_path, width=width)
    new_el = new_p._element
    new_el.getparent().remove(new_el)
    target_el.addnext(new_el)
    return new_el


# Буцаасан дарааллаар шинэчилбэл индекс нүүхгүй
for idx in reversed(range(len(image_locations))):
    para_idx = image_locations[idx]
    fig_name = fig_files[idx]
    img_path = os.path.join(FIG_DIR, fig_name)
    print(f"  Replacing [{para_idx}] -> {fig_name}")

    old_p = d.paragraphs[para_idx]
    old_el = old_p._element
    # Шинийг өмнө нь оруулах
    insert_picture_after(old_el, img_path, FIG_WIDTH)
    # Хуучныг устгах
    old_el.getparent().remove(old_el)

# ==================== 3) F5-г оруулах ================================
# Зураг 5 caption (194)-ийн өмнө оруулна
# Refetch paras
paras = list(d.paragraphs)
f5_caption_idx = None
for i, p in enumerate(paras):
    if p.text.strip().startswith('Зураг 5'):
        f5_caption_idx = i
        break

if f5_caption_idx is None:
    raise RuntimeError('Зураг 5 caption not found')

print(f"Зураг 5 caption: [{f5_caption_idx}]")

# Caption-ийн өмнөх paragraph
f5_caption_el = paras[f5_caption_idx]._element
prev_el = f5_caption_el.getprevious()

# Шинэ зураг paragraph үүсгэх
new_p = d.add_paragraph()
run = new_p.add_run()
run.add_picture(os.path.join(FIG_DIR, 'f5_ols_vs_iv.png'), width=FIG_WIDTH)
new_el = new_p._element
new_el.getparent().remove(new_el)
# Зураг 5 caption-ийн өмнө оруулах
f5_caption_el.addprevious(new_el)

print("  F5 inserted")

# ==================== 4) Товчилсон үгсийн хүснэгтийг шинэчлэх ==============

# Table 1 нь товчилсон үгсийн хүснэгт
abbrev_table = d.tables[1]

# Шинэ мөрүүд
new_rows = [
    ('Товчлол', 'Англи нэр', 'Монгол нэр'),                                # header
    ('ЭХБК', 'Ordinary Least Squares (OLS)',
     'Энгийн хамгийн бага квадратын арга'),
    ('ХШХБК', 'Two-Stage Least Squares (2SLS)',
     'Хоёр шатлалт хамгийн бага квадратын арга'),
    ('ХХБР', 'IV Threshold Regression (IVTR)',
     'Хэрэгсэл хувьсагчтай босго утгат регресс'),
    ('ӨНЭЗС', 'Household Socio-Economic Survey (HSES)',
     'Өрхийн нийгэм, эдийн засгийн судалгаа'),
    ('ҮСХ', 'National Statistics Office',
     'Үндэсний Статистикийн Хороо'),
    ('ЕБС', 'General Education School',
     'Ерөнхий боловсролын сургууль'),
    ('МСҮТ', 'Vocational Education Training Centre',
     'Мэргэжлийн сургалт, үйлдвэрлэлийн төв'),
    ('ЭЕШ', 'University Entrance Examination',
     'Элсэлтийн ерөнхий шалгалт'),
    ('JEL', 'Journal of Economic Literature',
     'Эдийн засгийн уран зохиолын сэтгүүлийн ангилал'),
    ('APA', 'American Psychological Association',
     'Америкийн Сэтгэл Судлалын Холбооны ишлэлийн стандарт'),
]

# Хуучин мөрүүдийг устгах (header хадгална)
while len(abbrev_table.rows) > 1:
    row_el = abbrev_table.rows[-1]._element
    row_el.getparent().remove(row_el)

# Header-г шинэчлэх
header_cells = abbrev_table.rows[0].cells
for i, text in enumerate(new_rows[0]):
    cell = header_cells[i]
    # Cell-ийн агуулгыг цэвэрлээд шинэ text оруулах
    for para in cell.paragraphs:
        if para.runs:
            para.runs[0].text = ''
            for r in para.runs[1:]:
                r.text = ''
    cell.paragraphs[0].add_run(text).bold = True

# Шинэ data мөрүүдийг нэмэх
for row_data in new_rows[1:]:
    new_row = abbrev_table.add_row()
    for i, text in enumerate(row_data):
        cell = new_row.cells[i]
        # add_run-ээр шинэ текст
        cell.paragraphs[0].text = text

d.save(dst)
print(f"\nSaved: {dst}")

# ==================== Баталгаажуулалт =================================
d2 = Document(dst)

# Зургийн тоо
total_images = 0
image_paras = []
for i, p in enumerate(d2.paragraphs):
    if p._element.findall('.//' + qn('w:drawing')):
        total_images += 1
        image_paras.append(i)
print(f"\nНийт зураг: {total_images} (5 байх ёстой)")
print(f"Зурагтай parа: {image_paras}")

# Товчилсон үгсийн хүснэгт
print("\n=== Шинэчилсэн товчилсон үгсийн хүснэгт ===")
t = d2.tables[1]
for ri, row in enumerate(t.rows):
    cells = [c.text.strip() for c in row.cells]
    print(f"  {cells}")
