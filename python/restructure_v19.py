# -*- coding: utf-8 -*-
"""v19: 20 эх сурвалжийн стандартад нийцүүлэн БҮТЭЦ шинэчлэх.

Гол өөрчлөлт:
1. V бүлэг «ХЭЛЭЛЦҮҮЛЭГ»-ийг БҮРЭН УСТГАХ (Монгол уламжлалд байхгүй)
2. 5.2 «Өмнөх судалгаатай харьцуулалт» → IV бүлгийн шинэ 4.5-д шилжүүлэх
3. 5.3 «Судалгааны хязгаарлалт» → VI бүлгийн шинэ 6.3-д шилжүүлэх
4. VI бүлгийг V бүлэг болгож дугаарлах, дэд бүлгүүдийг нь 5.1, 5.2, 5.3 болгох
5. VI бүлгийн агуулгад үлдсэн OLS/IVTR/2SLS/Хойд Америкт хэсгүүдийг засах
"""
import sys
import io
import shutil
from docx import Document

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

src = 'outputs/paper/bagiin_ner_v18.docx'
dst = 'outputs/paper/bagiin_ner_v19.docx'
shutil.copy(src, dst)
d = Document(dst)


def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def insert_after(target_el, text, style='Normal'):
    new_p = d.add_paragraph(text, style=style)
    new_el = new_p._element
    new_el.getparent().remove(new_el)
    target_el.addnext(new_el)
    return new_el


# =============== Step 1: V бүлгийн агуулгыг хадгалах (шилжүүлэхийн тулд) ===

# 5.2-ын 3 догол мөрийг шинэ хувилбараар бэлтгэх (нэр томьёо засах)
harch_1 = (
    "Энэхүү судалгааны ЭХБК ба тогтмол нөлөөний үнэлгээ (6.8 хувь) "
    "нь Pastore (2010)-ийн 7–9 хувь, Дэлхийн банкны Montenegro, "
    "Patrinos (2014) нарын 10.1 хувийн тооцоотой ерөнхийдөө "
    "нийцэж байна. Гэхдээ уг өмнөх бүтээлүүд нь нэг нэгдсэн "
    "дундаж үнэлгээг гаргаж, босгоны хоёр талын ялгааг далдалж "
    "харуулсан. Иймд бодлогын тайланд Монгол Улсад боловсролын "
    "өгөөж дундаж 8 хувь орчим гэсэн нэг тооны дүр зураг бий "
    "болгож, бүх түвшний боловсролд ижил үр нөлөөтэй гэсэн "
    "төөрөгдүүлсэн ойлголтыг үүсгэсэн."
)

harch_2 = (
    "Энэхүү судалгааны ХХБР үр дүнг Pastore (2010)-ийн жигнэсэн "
    "дундаж тооцооллоор харвал хэрэв түүврийн 40 хувь нь 13 "
    "жилээс доош, 60 хувь нь 13 жилээс дээш боловсролтой байсан "
    "гэж үзэхэд жигнэсэн дундаж нь 0.4 × 5.5 + 0.6 × 17.9 ≈ 12.9 "
    "хувь байх ёстой. Энэ тоо нь ЭХБК-д хязгаарлагдсан Pastore "
    "(2010)-ийн 7–9 хувь үнэлгээнээс 40 хувиар өндөр байх "
    "байсан. Зөрүү нь ЭХБК-ийн доош чиглэсэн сулруулагч хазайлт "
    "болон босгыг огт авч үзээгүйгээс үүдсэн болох нь харагдаж "
    "байна."
)

harch_3 = (
    "Олон улсын хүрээнд Psacharopoulos, Patrinos (2018) нарын "
    "139 орныг хамарсан мета-шинжилгээнд боловсролын өгөөжийн "
    "дэлхийн дундаж жилд 9.0 хувь, дээд боловсролынх 14.6 хувь "
    "гэж тогтоосон. Манай судалгааны ХХБР-ын дээд регимийн 17.9 "
    "хувь нь энэхүү олон улсын жишгээс давж буй нь хөгжиж буй "
    "орнуудын хөдөлмөрийн зах зээлд дипломын ховор бэлэнтэй "
    "харилцан хамааралтай болно. Энэ нь Duflo (2001)-ийн "
    "таамагласнаар дипломын хомсдлын үед итгэмжлэлийн үнэ цэнэ "
    "харьцангуй өндөр байдагтай нийцэж байна."
)

# 5.3-ын 2 догол мөрийг шинэ хувилбараар (хязгаарлалт, цаашдын чиглэл)
hyzaar_1 = (
    "Энэхүү судалгаа нь хэд хэдэн хязгаарлалттай болохыг ил "
    "тод хүлээн зөвшөөрөх шаардлагатай. Нэгдүгээрт, шинжилгээ "
    "нь зөвхөн цалин хөлстэй ажиллагчдын түүвэрт үндэслэсэн тул "
    "өөрийн бизнес эрхлэгчид болон албан бус салбарын ажиллагчдыг "
    "хамаараагүй. Энэ нь β-г бага зэрэг дээш чиглэсэн хазайлтад "
    "оруулах магадлалтай — учир нь өндөр боловсролтой ч өөрийн "
    "бизнесээр бага орлоготой байгаа ажиллагчид түүвэрт орох "
    "боломжгүй. Хекманы засвар нь уг хазайлтыг статистикийн "
    "хувьд ач холбогдолгүй гэж баталсан хэдий ч бүрэн арилгаж "
    "чадаагүй."
)

hyzaar_2 = (
    "Хоёрдугаарт, 2016, 2018 оны ӨНЭЗС-ийн давалгаанд төрсөн "
    "аймгийн мэдээлэл оруулаагүй байсан тул үндсэн хэрэгсэл "
    "хувьсагчтай шинжилгээнд оруулж чадаагүй. Энэ нь түүврийн "
    "хэмжээг багасгаж стандарт алдааг өсгөж болох хэдий ч "
    "коэффициентэд хазайлт үүсгэхгүй болохыг бат бөх байдлын "
    "шинжилгээгээр баталсан. Гуравдугаарт, боловсролын чанарын "
    "хувьсагч (шалгалтын оноо, сургуулийн нэр, мэргэжлийн ангилал) "
    "нь ӨНЭЗС-д оруулагдаагүй тул зөвхөн боловсролын жилийн "
    "тоог ашиглах шаардлагатай болсон. Ур чадварын далд хазайлтыг "
    "хэрэгсэл хувьсагч хэсэгчлэн засч буй ч чанарын дотоод "
    "ялгааг бүрэн барих боломжгүй. Цаашдын судалгаанд ӨНЭЗС-д "
    "боловсролын чанарын үзүүлэлтийг нэмж оруулах нь эдгээр "
    "хязгаарлалтыг арилгах чухал алхам болно."
)

# =============== Step 2: V бүлгийг олж устгах ======================
start_V = None
end_V = None
paras = list(d.paragraphs)
for i, p in enumerate(paras):
    if p.style.name == 'Heading 1':
        if 'V БҮЛЭГ' in p.text and 'VI' not in p.text:
            start_V = i
        elif start_V is not None and 'VI БҮЛЭГ' in p.text:
            end_V = i
            break

print(f"V бүлэг: {start_V}..{end_V}")

# V бүлгийн бүх догол мөрийг хадгалж аваад устгах
to_delete = [paras[i] for i in range(start_V, end_V)]
for p in to_delete:
    p._element.getparent().remove(p._element)

# =============== Step 3: IV бүлгийн 4.5-г 4.6 болгож өөрчлөх =========
# Refetch
paras = list(d.paragraphs)
i_45 = None
for i, p in enumerate(paras):
    if (p.style.name == 'Heading 2' and
            p.text.strip().startswith('4.5.')):
        i_45 = i
        break

if i_45 is None:
    raise RuntimeError("4.5 heading not found")

old_45_text = paras[i_45].text
# 4.5 → 4.6
new_46_text = old_45_text.replace('4.5.', '4.6.')
set_text(paras[i_45], new_46_text)

# =============== Step 4: 4.5 шинэ heading + харьцуулалтын 3 догол =====
# 4.5 шинэ = "Олон улсын туршлагатай харьцуулалт"
# 4.4 (Бат бөх байдал)-ийн сүүлийн paragraph-ийн дараа оруулах

# 4.4 хэсэгт [201, 202] prose paragraphs байгаа, тэдгээрийн сүүлийн нэг нь 202
# Refetch
paras = list(d.paragraphs)
i_44 = None
for i, p in enumerate(paras):
    if (p.style.name == 'Heading 2' and
            p.text.strip().startswith('4.4.')):
        i_44 = i
        break

# 4.4-н сүүлийн paragraph-ийг олох — дараагийн Heading 2 (4.6) хүртэл
i_46_now = None
for i, p in enumerate(paras):
    if (p.style.name == 'Heading 2' and
            p.text.strip().startswith('4.6.')):
        i_46_now = i
        break

# 4.4-н сүүлийн прозыг олох
last_44 = i_46_now - 1
while paras[last_44].style.name != 'Normal' or not paras[last_44].text.strip():
    last_44 -= 1

# 4.4-н сүүлийн prose-ын дараа 4.5 heading ба 3 харьцуулалт paragraph оруулна
target = paras[last_44]._element
new45_head = insert_after(target, "4.5. Олон улсын туршлагатай харьцуулалт", 'Heading 2')
h1 = insert_after(new45_head, harch_1, 'Normal')
h2 = insert_after(h1, harch_2, 'Normal')
h3 = insert_after(h2, harch_3, 'Normal')

# =============== Step 5: VI бүлгийг V бүлэг болгох, дэд бүлгүүдийг renumber =====
paras = list(d.paragraphs)
for p in paras:
    t = p.text.strip()
    if p.style.name == 'Heading 1' and 'VI БҮЛЭГ' in t:
        new_text = t.replace('VI БҮЛЭГ', 'V БҮЛЭГ')
        set_text(p, new_text)
    elif p.style.name == 'Heading 2' and t.startswith('6.1.'):
        set_text(p, t.replace('6.1.', '5.1.'))
    elif p.style.name == 'Heading 2' and t.startswith('6.2.'):
        set_text(p, t.replace('6.2.', '5.2.'))

# =============== Step 6: VI (одоо V)-ийн агуулгад үлдсэн OLS/IVTR/2SLS засах ==

def fix_runs(p, pairs):
    for run in p.runs:
        orig = run.text
        new = orig
        for old, repl in pairs:
            new = new.replace(old, repl)
        if new != orig:
            run.text = new


translations = [
    ('OLS-ийн 6.8%-ийн дундаж үнэлгээ', 'ЭХБК-ийн 6.8 хувийн дундаж үнэлгээ'),
    ('OLS-ийн 6.8%-аас 2SLS-ийн 11.3% хүртэлх',
     'ЭХБК-ийн 6.8 хувиас ХШХБК-ийн 11.3 хувь хүртэлх'),
    ('Хойд Америкт', 'Америкийн Нэгдсэн Улсад'),
    ('IVTR шинжилгээний', 'ХХБР шинжилгээний'),
    ('IVTR-ын тоон', 'ХХБР-ийн тоон'),
    ('IVTR-ын', 'ХХБР-ийн'),
    ('OLS-ийн', 'ЭХБК-ийн'),
    ('"дипломын нэмэгдэл" онолыг',
     'дипломын нэмэгдэл онолыг'),
    ('"дипломын нэмэгдэл" үзэгдэл оршин',
     'дипломын нэмэгдэл үзэгдэл оршин'),
]

# Хаана ч байж болно V бүлэг
for p in d.paragraphs:
    fix_runs(p, translations)

# =============== Step 7: V бүлгийн 5.3 шинэ subsection нэмэх ==========
# V бүлгийн эцэст 5.3 хязгаарлалт heading + 2 догол оруулна
paras = list(d.paragraphs)

# V бүлгийн төгсгөл олох: 5.2 дэд бүлгийн дараах сүүлийн прозыг олно
i_v_start = None
for i, p in enumerate(paras):
    if p.style.name == 'Heading 1' and ('V БҮЛЭГ' in p.text and
                                          'VI' not in p.text):
        i_v_start = i
        break

# Одоогийн V бүлгийн сүүлийн прозыг олох (НОМ ЗҮЙ-гээс өмнө)
i_end = None
for i in range(i_v_start + 1, len(paras)):
    p = paras[i]
    if p.style.name == 'Heading 1':
        i_end = i
        break

if i_end is None:
    i_end = len(paras)

# сүүлийн прозыг олох
last_prose = i_end - 1
while last_prose > i_v_start and (
        paras[last_prose].style.name != 'Normal' or
        not paras[last_prose].text.strip()):
    last_prose -= 1

target = paras[last_prose]._element
h53 = insert_after(target, "5.3. Судалгааны хязгаарлалт, цаашдын судалгааны чиглэл",
                    'Heading 2')
l1 = insert_after(h53, hyzaar_1, 'Normal')
l2 = insert_after(l1, hyzaar_2, 'Normal')

d.save(dst)
print("Saved:", dst)

# =============== Баталгаажуулалт ==============================
d2 = Document(dst)
print("\n=== БҮТЭЦ ===")
for i, p in enumerate(d2.paragraphs):
    if p.style.name in ('Heading 1', 'Heading 2'):
        t = p.text.strip()
        if any(k in t for k in ['БҮЛЭГ', 'АГУУЛГА', 'ХУРААНГУЙ',
                                  'ОРШИЛ', 'НОМ ЗҮЙ', 'ХАВСРАЛТ',
                                  'жагсаалт', 'ҮГС', 'тайлбар']):
            print(f"  [{i}] {p.style.name}: {t}")
        elif t and t[0].isdigit() and '.' in t[:4]:
            print(f"    [{i}] {p.style.name}: {t}")

# Foreign terms
bad = ['OLS', 'IVTR', '2SLS', 'фиксэлсэн', 'робастнес',
       'концентрацилагдсан', 'Хойд Америкт']
print("\n=== Гадаад үг ===")
for i, p in enumerate(d2.paragraphs):
    if i < 50:
        continue
    for b in bad:
        if b in p.text:
            print(f"  [{i}] {b!r}: {p.text[:100]}")
