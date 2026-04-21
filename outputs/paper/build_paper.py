"""
Build the SEZIS Econometrics Olympiad paper with proper hand-crafted
layout:  borders on tables, caption top-left, source bottom-right,
dot-leader table/figure lists, page breaks between major sections.

Starts from cover.docx (already has correct margins 25/20/30/20 and
team info) and appends everything else programmatically.
"""
import sys, shutil
from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap

sys.stdout.reconfigure(encoding='utf-8')

OUT_DIR = Path(__file__).parent
COVER = OUT_DIR / "cover.docx"
FINAL = OUT_DIR / "bagiin_ner_v5.docx"
FIGURES = OUT_DIR.parent.parent / "inputs" / "figures"

PAGE_CONTENT_WIDTH_MM = 210 - 30 - 20  # 160 mm


# ─────────────────────────────────────────────────────────────────────────────
# Equations loader: parse OMML blocks from Pandoc-generated equations.docx
# ─────────────────────────────────────────────────────────────────────────────

def load_equations():
    """Read equations.docx and return dict: name -> oMathPara XML element."""
    eq_path = OUT_DIR / "equations.docx"
    if not eq_path.exists():
        print(f"Warning: {eq_path} not found. Equations will be skipped.")
        return {}
    eq_doc = Document(str(eq_path))
    # Pandoc generates paragraphs in order: label, empty, math, label, empty, math, ...
    result = {}
    current_label = None
    for p in eq_doc.paragraphs:
        txt = p.text.strip()
        if txt.startswith("EQ_"):
            current_label = txt
        # Check if this paragraph contains an oMath or oMathPara
        xml = p._element
        maths = xml.findall('.//' + qn('m:oMathPara'))
        if not maths:
            maths = xml.findall('.//' + qn('m:oMath'))
        for m in maths:
            if current_label and current_label not in result:
                result[current_label] = deepcopy(m)
                current_label = None
                break
    print(f"Loaded {len(result)} equations: {sorted(result.keys())}")
    return result


def add_equation(doc, omath_element, number=None):
    """Insert an oMathPara/oMath element as a new paragraph.
    If number is given, adds a right-flush equation number like (1), (2)."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5
    if number is not None:
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.tab_stops.add_tab_stop(Cm(8.0), WD_TAB_ALIGNMENT.CENTER)
        pf.tab_stops.add_tab_stop(Cm(16.0), WD_TAB_ALIGNMENT.RIGHT)
        r0 = p.add_run('\t')
        set_run_props(r0, size=12)
        p._element.append(deepcopy(omath_element))
        r1 = p.add_run(f'\t({number})')
        set_run_props(r1, size=12)
    else:
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p._element.append(deepcopy(omath_element))
    return p


# ─────────────────────────────────────────────────────────────────────────────
# Bookmarks and PAGEREF fields
# ─────────────────────────────────────────────────────────────────────────────

_BOOKMARK_ID = [100]  # mutable counter

def add_bookmark(paragraph, name):
    """Wrap paragraph content with bookmark start/end."""
    bid = _BOOKMARK_ID[0]
    _BOOKMARK_ID[0] += 1
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(bid))
    start.set(qn('w:name'), name)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(bid))
    # Insert bookmark start at beginning, end at end
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is not None:
        pPr.addnext(start)
    else:
        paragraph._element.insert(0, start)
    paragraph._element.append(end)


def add_page_ref(paragraph, bookmark_name, size=12, bold=False):
    """Insert a PAGEREF field referencing the given bookmark."""
    r = paragraph.add_run()
    set_run_props(r, size=size, bold=bold)
    # <w:fldChar begin>
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    r._element.append(fld_begin)
    # <w:instrText>
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = f' PAGEREF {bookmark_name} \\h '
    r._element.append(instr)
    # <w:fldChar separate>
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    r._element.append(fld_sep)
    # placeholder (shown until user presses F9)
    t = OxmlElement('w:t')
    t.text = '—'
    r._element.append(t)
    # <w:fldChar end>
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r._element.append(fld_end)


def add_toc_field(paragraph, levels="1-2"):
    """Insert a Word TOC field that auto-generates on F9/open."""
    r = paragraph.add_run()
    set_run_props(r, size=12)
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    r._element.append(fld_begin)
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = f' TOC \\o "{levels}" \\h \\z \\u '
    r._element.append(instr)
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    r._element.append(fld_sep)
    t = OxmlElement('w:t')
    t.text = 'Агуулгыг шинэчлэхийн тулд F9 товч дарна уу.'
    r._element.append(t)
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r._element.append(fld_end)


def set_update_fields_on_open(doc):
    """Tell Word to auto-update fields when the document opens."""
    settings = doc.settings.element
    existing = settings.find(qn('w:updateFields'))
    if existing is not None:
        settings.remove(existing)
    uf = OxmlElement('w:updateFields')
    uf.set(qn('w:val'), 'true')
    settings.append(uf)

# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def set_run_props(run, font='Times New Roman', size=12, bold=False, italic=False,
                  color=None, lang='mn-MN'):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    # East-Asian font binding for Mongolian chars
    rFonts = run._element.rPr.rFonts if run._element.rPr is not None and run._element.rPr.find(qn('w:rFonts')) is not None else None
    if rFonts is None:
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)
    rFonts.set(qn('w:cs'), font)
    rFonts.set(qn('w:eastAsia'), font)


def set_paragraph_spacing(p, line_spacing=1.5, before=6, after=6,
                          alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    p.alignment = alignment


def add_paragraph(doc, text='', bold=False, italic=False, size=12,
                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_first=True,
                  line_spacing=1.5, before=6, after=6):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, line_spacing=line_spacing, before=before, after=after,
                          alignment=alignment)
    if indent_first and alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
        p.paragraph_format.first_line_indent = Cm(1.0)
    if text:
        run = p.add_run(text)
        set_run_props(run, size=size, bold=bold, italic=italic)
    return p


def add_mixed(doc, parts, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_first=True,
              size=12, line_spacing=1.5, before=6, after=6):
    """parts = list of (text, props_dict)."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, line_spacing=line_spacing, before=before, after=after,
                          alignment=alignment)
    if indent_first and alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
        p.paragraph_format.first_line_indent = Cm(1.0)
    for text, props in parts:
        run = p.add_run(text)
        set_run_props(run, size=props.get('size', size),
                      bold=props.get('bold', False),
                      italic=props.get('italic', False))
    return p


def add_heading(doc, text, level=1, before_break=False):
    """
    level 1 = chapter title, 14pt bold LEFT aligned, uppercase
    level 2 = section heading, 12pt bold left
    level 3 = subsection heading, 12pt bold italic left
    """
    if before_break:
        add_page_break(doc)
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if level == 1:
        pf.space_before = Pt(18)
        pf.space_after = Pt(12)
        run = p.add_run(text.upper())
        set_run_props(run, size=14, bold=True)
        pPr = p._element.get_or_add_pPr()
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), 'Heading1')
        pPr.insert(0, pStyle)
    elif level == 2:
        pf.space_before = Pt(12)
        pf.space_after = Pt(6)
        run = p.add_run(text)
        set_run_props(run, size=12, bold=True)
        pPr = p._element.get_or_add_pPr()
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), 'Heading2')
        pPr.insert(0, pStyle)
    else:  # level 3
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        run = p.add_run(text)
        set_run_props(run, size=12, bold=True, italic=True)
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    return p


# ─────────────────────────────────────────────────────────────────────────────
# Table helpers
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_borders(cell, sz=4, color='000000'):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for b in ['top', 'left', 'bottom', 'right']:
        tag = qn(f'w:{b}')
        existing = tcBorders.find(tag)
        if existing is not None:
            tcBorders.remove(existing)
        border = OxmlElement(f'w:{b}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(sz))
        border.set(qn('w:color'), color)
        tcBorders.append(border)


def set_cell_shade(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def write_cell(cell, text, bold=False, italic=False, size=10, align='left'):
    # Remove default paragraph
    cell.text = ''
    p = cell.paragraphs[0]
    pf = p.paragraph_format
    pf.line_spacing = 1.15
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.first_line_indent = Cm(0)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_props(run, size=size, bold=bold, italic=italic)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table_caption(doc, number, title):
    """Caption placed ABOVE table, left-aligned. Format: 'Хүснэгт N.M  Title'.
    Adds a bookmark so PAGEREF fields can link back to this page."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p.add_run(f"Хүснэгт {number}  ")
    set_run_props(r1, size=11, bold=True)
    r2 = p.add_run(title)
    set_run_props(r2, size=11, italic=True)
    # Bookmark name must not contain '.' — replace with '_'
    bm_name = f"Tbl_{number.replace('.', '_')}"
    add_bookmark(p, bm_name)


def add_table_source(doc, text="Эх сурвалж: Оюутны тооцоолол"):
    """Source placed BELOW table, right-aligned, 10pt italic."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(12)
    pf.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    set_run_props(run, size=10, italic=True)


def add_figure_caption(doc, number, title):
    """Caption placed BELOW figure, centered. Format: 'Зураг N.M  Title'.
    Adds a bookmark so PAGEREF fields can link back to this page."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f"Зураг {number}  ")
    set_run_props(r1, size=11, bold=True)
    r2 = p.add_run(title)
    set_run_props(r2, size=11, italic=True)
    bm_name = f"Fig_{number.replace('.', '_')}"
    add_bookmark(p, bm_name)


def add_figure_source(doc, text="Эх сурвалж: Оюутны тооцоолол"):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(12)
    pf.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    set_run_props(run, size=10, italic=True)


def build_table(doc, headers, rows, col_widths_cm=None, header_shade='DDDDDD',
                align='center', first_col_left=True):
    """Build a bordered Word table.
       rows can contain dicts with special 'note' key to render as merged note row.
    """
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Column widths
    if col_widths_cm is not None:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Cm(col_widths_cm[i])

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        write_cell(cell, h, bold=True, size=10, align='center')
        set_cell_shade(cell, header_shade)
        set_cell_borders(cell)

    # Data rows
    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row):
            cell = tr.cells[c_idx]
            # First column left-aligned if requested
            if c_idx == 0 and first_col_left:
                write_cell(cell, str(val), size=10, align='left')
            else:
                write_cell(cell, str(val), size=10, align=align)
            set_cell_borders(cell)
    return table


def add_image_centered(doc, image_path, width_cm=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    return p


# ─────────────────────────────────────────────────────────────────────────────
# List pages: Хүснэгтэн/Зурган жагсаалт with dot-leader tabs
# ─────────────────────────────────────────────────────────────────────────────

def add_list_entry(doc, label, title, bookmark=None):
    """Table/Figure list entry with dot leader + PAGEREF field to bookmark."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15
    right_tab_cm = 16.0
    pf.tab_stops.add_tab_stop(Cm(right_tab_cm), WD_TAB_ALIGNMENT.RIGHT,
                              WD_TAB_LEADER.DOTS)
    r1 = p.add_run(f"{label}. ")
    set_run_props(r1, size=12, bold=True)
    r2 = p.add_run(title)
    set_run_props(r2, size=12)
    r3 = p.add_run('\t')
    set_run_props(r3, size=12)
    if bookmark:
        add_page_ref(p, bookmark)
    else:
        r4 = p.add_run('—')
        set_run_props(r4, size=12)
    return p


# ─────────────────────────────────────────────────────────────────────────────
# Tables and figures registry (for internal dogga)
# ─────────────────────────────────────────────────────────────────────────────

TABLES = [
    ("1", "Шинжилгээний түүврийн үндсэн үзүүлэлт"),
    ("2", "Түүврийн үндсэн үзүүлэлт ӨНЭЗС-ийн давалгаа тус бүрээр"),
    ("3", "Энгийн OLS ба фиксэлсэн нөлөөтэй үнэлгээний үр дүн"),
    ("4", "Хэрэгсэл хувьсагчийн (IV) үнэлгээний үр дүн"),
    ("5", "Хэрэгсэл хувьсагчтай босготой регресс (IVTR) — үндсэн үр дүн"),
    ("6", "Үндсэн үр дүнгийн робастнес шинжилгээ"),
    ("7", "Эконометрик загваруудын нэгдсэн дүгнэлт"),
]

FIGURES_LIST = [
    ("1", "Боловсролын жилийн тархалт ӨНЭЗС-ийн давалгаа бүрээр",
     FIGURES / "f1_education_distribution.png"),
    ("2", "Боловсролын жил ба логаритмчилсан цалингийн хамаарал (ӨНЭЗС 2024)",
     FIGURES / "f2_wage_education_scatter.png"),
    ("3", "Босгоны сүлжээн хайлтын төвлөрсөн алдааны квадратын нийлбэр",
     FIGURES / "f3_threshold_profile.png"),
    ("4", "Босгоор тусгаарлагдсан боловсролын IV өгөөж (Caner–Hansen IVTR)",
     FIGURES / "f4_regime_slopes.png"),
    ("5", "OLS ба IV үнэлгээний давалгаа хоорондын харьцуулалт",
     FIGURES / "f5_ols_vs_iv.png"),
]

ABBREV = [
    ("OLS", "Ordinary Least Squares", "Хамгийн бага квадратын энгийн арга"),
    ("2SLS", "Two-Stage Least Squares", "Хоёр шатлалт хамгийн бага квадратын арга"),
    ("IV", "Instrumental Variable", "Хэрэгсэл хувьсагч"),
    ("IVTR", "Instrumental Variable Threshold Regression",
     "Хэрэгсэл хувьсагчтай босготой регресс"),
    ("FE", "Fixed Effects", "Фиксэлсэн нөлөөний загвар"),
    ("RE", "Random Effects", "Санамсаргүй нөлөөний загвар"),
    ("SE", "Standard Error", "Стандарт алдаа"),
    ("CI", "Confidence Interval", "Итгэлийн интервал"),
    ("SSR", "Sum of Squared Residuals", "Алдааны квадратын нийлбэр"),
    ("ӨНЭЗС", "Household Socio-Economic Survey (HSES)",
     "Өрхийн нийгэм эдийн засгийн судалгаа"),
    ("ҮСХ", "National Statistics Office", "Үндэсний статистикийн хороо"),
    ("ЕБС", "General Education School", "Ерөнхий боловсролын сургууль"),
    ("МСҮТ", "Vocational Education Training Centre",
     "Мэргэжлийн сургалт үйлдвэрлэлийн төв"),
    ("AR", "Anderson–Rubin test", "Андерсон–Рубиний тест"),
    ("KP", "Kleibergen–Paap", "Клейберген–Паапын статистик"),
    ("APA", "American Psychological Association",
     "Америкийн сэтгэл судлалын холбооны ишлэлийн стандарт"),
]


# ─────────────────────────────────────────────────────────────────────────────
# Build
# ─────────────────────────────────────────────────────────────────────────────

def main():
    print(f"Copying cover -> final: {FINAL.name}")
    shutil.copy(str(COVER), str(FINAL))
    doc = Document(str(FINAL))
    EQ = load_equations()

    # Ensure Normal style carries Times New Roman 12 / 1.5 with Mongolian lang
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_before = Pt(6)
    style.paragraph_format.space_after = Pt(6)

    # Page break after cover
    add_page_break(doc)

    # ─── АГУУЛГА (Table of Contents) ─ Word TOC field ───────────────────────
    add_heading(doc, "Агуулга", level=1)
    p_toc = doc.add_paragraph()
    set_paragraph_spacing(p_toc, line_spacing=1.15, before=0, after=0,
                          alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_toc_field(p_toc, levels="1-1")

    # ─── ХҮСНЭГТЭН МЭДЭЭЛЛИЙН ЖАГСААЛТ ───────────────────────────────────────
    add_heading(doc, "Хүснэгтэн мэдээллийн жагсаалт", level=1, before_break=True)
    for num, title in TABLES:
        add_list_entry(doc, f"Хүснэгт {num}", title,
                       bookmark=f"Tbl_{num.replace('.', '_')}")

    # ─── ЗУРГАН МЭДЭЭЛЛИЙН ЖАГСААЛТ ──────────────────────────────────────────
    add_heading(doc, "Зурган мэдээллийн жагсаалт", level=1, before_break=True)
    for num, title, _path in FIGURES_LIST:
        add_list_entry(doc, f"Зураг {num}", title,
                       bookmark=f"Fig_{num.replace('.', '_')}")

    # ─── ТОВЧИЛСОН ҮГС ───────────────────────────────────────────────────────
    add_heading(doc, "Товчилсон үгс", level=1, before_break=True)
    add_paragraph(doc,
        "Дараах хүснэгтэд энэхүү тайланд ашигласан товчилсон нэр томьёонуудыг "
        "англи ба монгол утгын хамт жагсаасан болно.",
        indent_first=True)
    tbl = doc.add_table(rows=1 + len(ABBREV), cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    widths = [3.0, 6.5, 6.5]
    for row in tbl.rows:
        for i, c in enumerate(row.cells):
            c.width = Cm(widths[i])
    headers = ["Товчилсон", "Англи нэр", "Монгол нэр"]
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        write_cell(cell, h, bold=True, size=10, align='center')
        set_cell_shade(cell, 'DDDDDD')
        set_cell_borders(cell)
    for r_idx, (abbr, en, mn) in enumerate(ABBREV):
        row = tbl.rows[r_idx + 1]
        write_cell(row.cells[0], abbr, bold=True, size=10, align='left')
        write_cell(row.cells[1], en, size=10, align='left')
        write_cell(row.cells[2], mn, size=10, align='left')
        for c in row.cells:
            set_cell_borders(c)

    # ─── ХУРААНГУЙ ───────────────────────────────────────────────────────────
    add_heading(doc, "Хураангуй", level=1, before_break=True)

    abstract_mn = [
        "Энэхүү судалгааны ажлаар Монгол Улс дахь хөдөлмөрийн зах зээлд "
        "боловсролын бодит өгөөжийг Caner ба Hansen (2004)-ийн хэрэгсэл "
        "хувьсагчтай босготой регрессийн (IVTR) аргаар үнэлэв. Өгөгдлийн "
        "сурвалжаар Үндэсний Статистикийн Хороо (ҮСХ)-ны Өрхийн нийгэм, "
        "эдийн засгийн судалгааны (ӨНЭЗС) 2016, 2018, 2020, 2021, 2024 "
        "оны таван давалгааг ашигласан бөгөөд нийт 272,096 хувь хүний "
        "бүртгэлээс 25–60 насны цалинтай ажилтан 49,366 хүнийг үндсэн "
        "шинжилгээний түүвэрт хамруулав. Уг судалгаа нь Caner–Hansen-ийн "
        "IVTR аргыг Монголын өрхийн микро өгөгдөлд анх удаа хэрэглэсэн "
        "эмпирик ажил юм.",

        "Шинжилгээг дөрвөн шатлалтай явуулав. Нэгдүгээрт, Минсерийн "
        "сонгодог цалингийн тэгшитгэлийг энгийн OLS аргаар үнэлж "
        "боловсролын нэмэлт нэг жилийн өгөөжийг 6.8% хэмээн тогтоов. "
        "Хоёрдугаарт, төрсөн когорт × аймаг × давалгаагаар 692 нүдтэй "
        "псевдо-панел үүсгэж, фиксэлсэн нөлөөний (FE) ба санамсаргүй "
        "нөлөөний (RE) загваруудыг үнэлсэн бөгөөд Хаусманы шалгуураар "
        "FE-ийг сонгов. Гуравдугаарт, боловсролын эндогений шинжийг "
        "засахын тулд (i) төрсөн аймгийн ангилал хувьсагч, (ii) "
        "Улаанбаатар хүртэлх зайн логарифм, (iii) сумын ерөнхий "
        "боловсролын сургуулийн (ЕБС) багш/сурагчийн харьцаа гэсэн "
        "гурван хэрэгсэл хувьсагчтайгаар 2SLS үнэлгээг хийж өгөөжийг "
        "11.3%-иар тооцов. Дөрөвдүгээрт, Caner–Hansen IVTR аргыг "
        "хэрэгжүүлж боловсролын жилийн хэмжээнд бүтцийн хугарлыг хайв.",

        "Судалгааны гол үр дүн нь Монголын хөдөлмөрийн зах зээлд "
        "боловсролын өгөөж нь 13 жилийн босго дээр мэдэгдэхүйц бүтцийн "
        "хугаралтай болохыг тогтоосон явдал юм. Босгоноос доош (educ ≤ 13) "
        "нэг жилийн өгөөж дунджаар 5.5%, босгоноос дээш (educ > 13) 17.9% "
        "буюу 3.3 дахин их байгаа нь Монголын хөдөлмөрийн зах зээлд "
        "дипломын нэмэгдэл эрс илэрч буйг харуулж байна. Бүтцийн хугарлын "
        "статистик ач холбогдлыг 1,000 удаагийн дахин түүвэрлэх шалгуурт "
        "(wild bootstrap) суурилсан SupWald статистикаар шалгахад p < 0.001 "
        "гарч, тохиолдлын шинжгүй гэсэн дүгнэлтэд хүрэв. Хекманы сонголтын "
        "засвар ба хяналтын тест (placebo) нь хэрэгсэл хувьсагчийн зөв "
        "үйлчилгээг баталгаажуулав.",

        "Судалгааны шинжлэх ухааны хувь нэмэр дараах дөрвөн чиглэлд "
        "оршино. Нэгдүгээрт, Caner–Hansen (2004)-ийн IVTR аргыг Монголын "
        "хөдөлмөрийн зах зээлд анх удаа хэрэглэсэн бөгөөд Даваажаргал, "
        "Цолмон (2019)-ийн макро түвшний Threshold SVAR шинжилгээтэй "
        "зэрэгцэн Монголын босго шинжилгээний эмпирик уламжлалыг "
        "өргөжүүлэв. Хоёрдугаарт, Card (1993)-ын газарзүйн ойртоцын "
        "хэрэгсэл хувьсагчийн стратегийг Монгол Улсын нутаг дэвсгэрийн "
        "онцлогт шилжүүлэн хэрэглэв. Гуравдугаарт, 1212.mn нээлттэй "
        "мэдээллийн сангаас сумын ЕБС-ийн багш/сурагчийн харьцааг "
        "татаж хэрэгсэл хувьсагч болгон ашиглав. Дөрөвдүгээрт, өгөөжийн "
        "босгон дээрх огцом үсрэлтийг эмпирикээр тогтоосон нь дээд "
        "боловсролын тэтгэлэг, ЕБС-ийн чанарыг сайжруулах, ӨНЭЗС-ийг "
        "шинэчлэх гурван чиглэлийн бодлогын саналын эмпирик үндэс болов.",
    ]
    for para in abstract_mn:
        add_paragraph(doc, para, indent_first=True)

    add_mixed(doc, [
        ("Түлхүүр үг: ", {'bold': True, 'size': 12}),
        ("боловсролын өгөөж, хэрэгсэл хувьсагч, босготой регресс, "
         "Caner-Hansen, Монгол, ӨНЭЗС.", {'size': 12})
    ], indent_first=False)

    add_heading(doc, "Abstract (English)", level=2)

    abstract_en = [
        "This paper estimates the causal return to schooling in Mongolia using an "
        "instrumental-variable threshold regression (IV-Threshold) approach. We "
        "combine five waves of the Household Socio-Economic Survey (HSES) 2016, "
        "2018, 2020, 2021 and 2024, pooling 272,096 individual records and "
        "retaining 49,366 wage-earning workers aged 25–60 for the main analysis.",

        "We proceed in four steps. A conventional Mincer OLS wage regression "
        "yields a 6.8% return per additional year of schooling. Fixed-effects "
        "and random-effects estimators on a pseudo-panel of 692 cohort × aimag "
        "× wave cells are compared via a Hausman test, which selects FE. To "
        "address the endogeneity of schooling, we construct three instruments "
        "— birth-aimag dummies, log distance to Ulaanbaatar, and the soum-"
        "level primary-school teacher-to-pupil ratio — and estimate 2SLS, "
        "obtaining a 12.0% return (first-stage F = 10.46). Finally, we "
        "implement the Caner and Hansen (2004) IV-threshold regression for "
        "the first time on Mongolian micro data.",

        "The central finding is a statistically significant structural break "
        "in the return to schooling at 13 years. Below the threshold the "
        "marginal return is 5.5% per year; above it the marginal return jumps "
        "to 17.9% per year — a 3.3-fold increase. A wild-bootstrap SupWald "
        "test with 1,000 replications rejects the no-threshold null at "
        "p < 0.001. Heckman selection correction and placebo tests confirm "
        "the validity of the IV design.",
    ]
    for para in abstract_en:
        add_paragraph(doc, para, indent_first=True)

    add_mixed(doc, [
        ("Keywords: ", {'bold': True, 'size': 12}),
        ("return to education, instrumental variable, threshold regression, "
         "Caner-Hansen, Mongolia, HSES.", {'size': 12})
    ], indent_first=False)

    # ─── ОРШИЛ ───────────────────────────────────────────────────────────────
    add_heading(doc, "Оршил", level=1, before_break=True)
    add_heading(doc, "Судалгааны сэдвийн үндэслэл", level=2)
    add_paragraph(doc,
        "Хүмүүн капиталын нэмэлт нэг жилийн боловсрол нь хөдөлмөрийн орлогыг "
        "хэрхэн нэмэгдүүлдэг вэ гэдэг асуулт эдийн засгийн ухааны хамгийн "
        "эртний бөгөөд үнэт суурь асуудлуудын нэг юм. Schultz (1961), Becker "
        "(1964) нарын үндэслэсэн хүмүүн капиталын онол нь боловсролыг хөрөнгө "
        "оруулалт хэмээн үзэж, түүнд зарцуулсан зардлыг ирээдүйн орлого "
        "хэлбэрээр нөхдөг гэж тайлбарладаг. Mincer (1974)-ийн цалингийн "
        "тэгшитгэл нь хожим олон арван жилийн турш энэ онолын эмпирик "
        "шалгалтын үндсэн хэрэгсэл болжээ. Psacharopoulos ба Patrinos (2018) "
        "нарын 139 орныг хамарсан мета-шинжилгээгээр боловсролын өгөөжийн "
        "дэлхийн дундаж жилд 9.0% орчим, дээд боловсролынх 14.6% болж "
        "байгааг тогтоосон.")
    add_paragraph(doc,
        "Монгол Улсын хувьд боловсролын өгөөжийн асуудал онцгой ач "
        "холбогдолтой болж байна. Нэгдүгээрт, 1995 онд 10 жилийн ерөнхий "
        "боловсролын тогтолцоог 11 жил болгож, 2008 оноос эхлэн 12 "
        "жилийн системд бүрэн шилжсэн шинэчлэлийн үр дагаварт боловсролын "
        "ахиу хувьсагчийн утга ажилтны үеүд дунд харилцан адилгүй байна. "
        "Хоёрдугаарт, 1990-ээд оноос хойш хувийн их, дээд сургуулиуд "
        "эрчимтэй өргөжин нэмэгдэж, дээд боловсролын хамралт 3 дахин "
        "өсч, 2024 оны байдлаар 25–34 насны иргэдийн 44% дээд боловсрол "
        "эзэмшжээ (ҮСХ, 2024). Гуравдугаарт, хөдөлмөрийн зах зээлд "
        "ажилгүйдлийн түвшин боловсролын түвшнээс эрс ялгаатай байна: "
        "бүрэн дунд боловсролтой иргэдийн 7.4% ажилгүй байхад дээд "
        "боловсролтой иргэдийн зөвхөн 5.6% нь ажилгүй байгаа нь "
        "боловсрол хөдөлмөрийн зах зээл дэх байр суурьтай хүчтэй "
        "холбоотой болохыг харуулна (Ай Ар Ай Эм, 2015).")
    add_paragraph(doc,
        "Монгол Улс дахь боловсролын өгөөжийн эмпирик судалгаа хомс "
        "хэвээр байна. Pastore (2010) 2007–2008 оны түүвэрт тулгуурлан "
        "залуучуудын боловсролын өгөөжийг 7–9% орчим хэмээн тогтоосон "
        "боловч уг бүтээлд шугаман бус байдал, эндогений хамаарал, "
        "хэрэгсэл хувьсагчийн засвар зэрэг орчин үеийн эконометрик "
        "сорилтуудыг бүрэн авч үзээгүй болно. Ай Ар Ай Эм ХХК (2015)-ийн "
        "Боловсрол, Соёл, Шинжлэх Ухаан, Спортын Яаманд хийсэн "
        "Хөдөлмөрийн зах зээлийн шинжилгээнд дээд боловсролын OLS өгөөжийг "
        "41% гэж үнэлсэн ч хэрэгсэл хувьсагч, босго шинжилгээ хийгдээгүй. "
        "Иймд Монголын өрхийн микро өгөгдөл дээр IV ба босго шинжилгээг "
        "хослуулан хэрэгжүүлэх нь өнөөгийн эмпирик уран зохиолын хоосрол "
        "юм.")

    add_heading(doc, "Судалгааны зорилго, зорилтууд", level=2)
    add_paragraph(doc,
        "Энэхүү судалгааны гол асуудлыг дараах гурван асуултаар томьёолов: "
        "(i) Монгол Улс дахь боловсролын нэмэлт нэг жилийн шалтгаант "
        "өгөөж хэр хэмжээтэй вэ?; (ii) уг өгөөж нь боловсролын жилийн "
        "хэмжээнээс хамаарсан бүтцийн босготой эсэх?; (iii) хэрэв босго "
        "оршдог бол тэр нь боловсролын хэддүгээр жил дээр тогтдог вэ?")
    add_paragraph(doc,
        "Судалгааны зорилго нь ӨНЭЗС-ийн өрхийн микро өгөгдөлд "
        "тулгуурлан Монгол Улс дахь боловсролын бодит өгөөжийн шугаман "
        "ба шугаман бус бүтцийг хэрэгсэл хувьсагчтай босготой регрессийн "
        "аргаар үнэлж, Монголын хөдөлмөрийн зах зээлд дипломын нэмэгдэл "
        "үзэгдэл (sheepskin effect) оршин буй эсэхэд эмпирик хариулт "
        "өгөх явдал юм.")
    add_paragraph(doc, "Энэ зорилгын хүрээнд дараах зорилтуудыг тавьсан:",
                  indent_first=False)
    objectives = [
        "Минсерийн цалингийн тэгшитгэлийг ӨНЭЗС-ийн таван давалгааны "
        "нэгтгэсэн өгөгдөл дээр OLS аргаар үнэлж, боловсролын өгөөжийн "
        "суурь үнэлгээг гарган авах;",
        "Төрсөн когорт × аймаг × давалгаагаар псевдо-панел үүсгэн "
        "фиксэлсэн болон санамсаргүй нөлөөний загваруудыг харьцуулж, "
        "Хаусманы шалгуураар тохирох загварыг сонгох;",
        "Боловсролын эндогений шинжийг засах зорилгоор гурван хэрэгсэл "
        "хувьсагчтайгаар (төрсөн аймаг, Улаанбаатар хүртэлх зай, сумын "
        "ЕБС-ийн багш/сурагчийн харьцаа) 2SLS үнэлгээг хийж, "
        "Клейберген–Паапын F статистикаар хэрэгсэл хувьсагчийн хүчийг "
        "шалгах;",
        "Caner ба Hansen (2004)-ийн хэрэгсэл хувьсагчтай босготой "
        "регрессийн аргачлалаар боловсролын жилийн хэмжээнд бүтцийн "
        "хугарал оршиж буй эсэхийг сүлжээн хайлтаар тогтоож, оновчтой "
        "босго γ*-ыг үнэлэх;",
        "Бүтцийн хугарлын статистик ач холбогдлыг 1,000 удаагийн дахин "
        "түүвэрлэх шалгууранд (wild bootstrap) тулгуурласан SupWald "
        "тестээр шалгах;",
        "Үр дүнг дэд түүврийн шинжилгээ, хяналтын тест (placebo), "
        "Хекманы сонголтын засвараар баталгаажуулж, тогтвортой байдлыг "
        "шалгах;",
        "Эмпирик үр дүнд тулгуурласан бодлогын саналыг боловсруулах.",
    ]
    for i, obj in enumerate(objectives, 1):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=3, after=3)
        p.paragraph_format.first_line_indent = Cm(1.0)
        run = p.add_run(f"{i}) {obj}")
        set_run_props(run, size=12)

    add_heading(doc, "Судалгааны шинэлэг тал", level=2)
    add_paragraph(doc,
        "Уг судалгаа нь дараах дөрвөн чиглэлд шинжлэх ухааны хувь нэмэр "
        "оруулав. Аргачлалын хувьд, Caner ба Hansen (2004)-ийн IVTR аргыг "
        "Монгол Улсын өрхийн микро өгөгдөлд анх удаа хэрэглэсэн бөгөөд уг "
        "алгоритмыг R хэл дээр гарын авлагаар хэрэгжүүлж, дахин "
        "хэрэглэгдэх код бэлтгэв. Онолын хувьд, Jaeger ба Page (1996)-ийн "
        "\"дипломын нэмэгдэл\" үзэгдэл (sheepskin effect) болон Spence "
        "(1973)-ийн дохионы онолын таамаглалыг Монголын хөдөлмөрийн зах "
        "зээл дээр эмпирикээр шалгасан анхны ажил юм. Өгөгдлийн хувьд, "
        "ӨНЭЗС-ийн таван давалгааны 272,096 хувь хүний бүртгэлийг нэгтгэн "
        "боловсруулж, 1212.mn нээлттэй мэдээллийн сангаас сумын ЕБС-ийн "
        "2000–2024 оны статистикийг татан хэрэгсэл хувьсагч болгон "
        "ашигласан нь шинэлэг. Бодлогын хувьд, босгон дээрх 3.3 дахин "
        "үсрэлтийн эмпирик нотолгоонд тулгуурлан Боловсрол, Шинжлэх "
        "ухааны яам, Хөдөлмөр, нийгмийн хамгааллын яамд зориулсан гурван "
        "шууд хэрэглээний саналыг боловсруулав.")

    add_heading(doc, "Судалгааны объект, хамрах хүрээ", level=2)
    add_paragraph(doc,
        "Судалгааны объект нь Монгол Улсын 25–60 насны цалин хөлстэй "
        "хөдөлмөр эрхэлж буй иргэд юм. Хамрах хүрээ нь ӨНЭЗС-ийн 2016, "
        "2018, 2020, 2021, 2024 оны таван давалгаа бөгөөд хугацааны "
        "хувьд найман жилийн түүхэн үе юм. Үндсэн шинжилгээнд "
        "2020+2021+2024 оны гурван давалгааны төрсөн аймгийн мэдээлэлтэй "
        "12,020 хүнийг ашигласан бол робастнес шинжилгээний үе шатанд "
        "2016+2018 оны давалгаануудыг нэмж оруулж нийт 49,366 хүний "
        "өгөгдлийг хамруулав. Ингэснээр шинжилгээний хамралтыг "
        "өргөжүүлж, үр дүнгийн тогтвортой байдлыг шалгах боломж бүрдэв.")

    # ─── НЭГДҮГЭЭР БҮЛЭГ. СУДЛАГДСАН БАЙДАЛ ──────────────────────────────────
    add_heading(doc, "I бүлэг. Судлагдсан байдал", level=1,
                before_break=True)

    add_heading(doc, "1.1. Боловсролын өгөөжийн сонгодог онол", level=2)
    add_paragraph(doc,
        "Боловсролын нэмэлт нэг жилийн цалингийн өгөөжийг эмпирикээр "
        "үнэлэх сонгодог хэлбэр нь Mincer (1974)-ийн цалингийн тэгшитгэл "
        "юм. Уг тэгшитгэл нь хүмүүн капиталын онолын (Schultz, 1961; "
        "Becker, 1964) практик хэрэглээ бөгөөд боловсролыг хөдөлмөрийн "
        "бүтээмжийг нэмэгдүүлэгч хөрөнгө оруулалт хэмээн үздэг. "
        "Монголбанкны судлаач Болдбаатар (2017) нар хүмүүн капиталын "
        "онолын үүднээс хувь хүний авъяас чадвар, эзэмшсэн мэргэжил "
        "нь эдийн засгийн өсөлтийн чухал хүчин зүйл болж, орчин үеийн "
        "мэдлэгт суурилсан эдийн засгийн хөгжлийн гол үндэс болдог "
        "болохыг тэмдэглэсэн байдаг.")
    add_paragraph(doc,
        "Минсерийн тэгшитгэл нь дараах үндсэн хэлбэртэй байна:")
    if 'EQ_MINCER_BASE' in EQ:
        add_equation(doc, EQ['EQ_MINCER_BASE'], number=1)
    add_paragraph(doc,
        "Энд lnw нь цалингийн логарифм, educ нь төгссөн боловсролын жил, "
        "exp нь ажлын туршлагын жил, ε нь алдаа гэсэн хувьсагчуудаас "
        "бүрдэнэ. β коэффициент нь боловсролын нэмэлт нэг жилийн "
        "цалинд үзүүлэх бодит өгөөжийг илэрхийлдэг. Psacharopoulos ба "
        "Patrinos (2018)-ийн 139 орныг хамарсан мета-шинжилгээгээр уг β "
        "коэффициентийн дэлхийн дундаж утга 9.0% бөгөөд, дээд боловсролын "
        "өгөөж нь дунджаар 14.6% гэж тогтоогдсон байдаг.")
    add_paragraph(doc,
        "Гэхдээ OLS үнэлгээ нь хоёр эндогений асуудалтай гэдэг нь урт "
        "хугацааны турш маргаан үүсгэсээр ирсэн. Нэгдүгээрт, хувь хүний "
        "боловсролын сонголт нь ажиглагдахгүй ур чадвар, авъяастай "
        "хамааралтай байдаг тул OLS-ийн β үнэлгээ дээш чиглэсэн хазайлтад "
        "орж, ур чадварын нөлөөг боловсролын нөлөө мэтээр хэмжиж болно. "
        "Хоёрдугаарт, боловсролын жилийн мэдээлэлд хэмжилтийн алдаа бий "
        "болох нь коэффициентийг доош чиглүүлсэн хэмжилтийн хазайлтад "
        "(attenuation bias) хүргэдэг. Эдгээр хоёр эсрэг чиглэлтэй "
        "хазайлт нь практикт хэсэгчлэн нөхөгддөг тул OLS үнэлгээ бодит "
        "параметрт хэр ойрт байгаа нь эхлэлтэй ч үл мэдэгдэх байна (Card, "
        "1999).")

    add_heading(doc, "1.2. Хэрэгсэл хувьсагчийн уламжлал", level=2)
    add_paragraph(doc,
        "OLS-ийн эндогений хазайлтыг засах хамгийн түгээмэл стратеги нь "
        "хэрэгсэл хувьсагч (IV) ашиглах арга юм. Card (1993) нь АНУ-д ойр "
        "коллежтэй нутагт өссөн байдал нь хувь хүний боловсролын шийдвэрт "
        "нөлөөлдөг боловч цалинд зөвхөн боловсролоор дамжин нөлөөлдөг "
        "гэсэн онолын үндэслэлээр коллежийн газарзүйн ойртоцыг хэрэгсэл "
        "хувьсагч болгон ашигласан. Уг судалгаанд IV-ийн 2SLS үнэлгээ нь "
        "OLS үнэлгээнээс 25–60%-иар өндөр гарсан нь боловсролын хэмжилтийн "
        "алдаанаас үүдэлтэй доош чиглэсэн хазайлт нь ур чадварын нөлөөнөөс "
        "давамгайлж байгааг илрүүлсэн чухал эмпирик нотолгоо байв.")
    add_paragraph(doc,
        "Duflo (2001) нь \"American Economic Review\" сэтгүүлд нийтлэгдсэн "
        "сонгодог бүтээлдээ Индонезийн INPRES сургуулийн бүтээх хөтөлбөрийг "
        "байгалийн туршилт болгон ашиглав. Уг бүтээлд 1973–1978 онд "
        "Индонез улс 1,000 хүүхдэд нэг шинэ ЕБС барьсан нь дундаж "
        "боловсролын жилийг 0.12–0.19 жилээр, цалинг 1.5–2.7%-иар "
        "нэмэгдүүлсэн болохыг тогтоосон. Үүнээс тооцох боловсролын "
        "өгөөж 6.8–10.6% бөгөөд, энэ нь хөгжиж буй орны сургуулийн "
        "бүтээн байгуулалт нь хүмүүн капиталын хуримтлалд шууд үр "
        "нөлөөтэй болохыг харуулж байна. Монголд ч үүнтэй төстэй "
        "сургуулийн нягтаршилд суурилсан стратеги ашиглах боломж "
        "байгаа боловч микро өгөгдөлд тулгуурласан эмпирик шалгалт "
        "одоог хүртэл хийгдээгүй болно.")

    add_heading(doc, "1.3. Босготой регрессийн уламжлал", level=2)
    add_paragraph(doc,
        "Босготой регресс нь өгөгдлийг тодорхой босго хувьсагчийн эргэн "
        "тойронд хоёр регимд хувааж, регим тус бүрт өөр параметр үнэлэх "
        "аргачлал юм. Hansen (2000) нь экзоген регрессортой тохиолдолд "
        "босгыг үнэлэх алгоритмыг боловсруулж, босгоны статистик ач "
        "холбогдлыг хэмжих SupWald шалгуурыг дэвшүүлжээ. SupWald "
        "статистик нь ердийн хи-квадратын хуваарилалтанд захирагддаггүй "
        "тул дахин түүвэрлэх шалгуур (bootstrap) өргөн хэрэглэгддэг.")
    add_paragraph(doc,
        "Hansen (2000)-ийн загвар зөвхөн экзоген регрессортой нөхцөлд "
        "хэрэглэгддэг учир боловсрол зэрэг эндогений шинжтэй хувьсагчид "
        "тохирохгүй байв. Энэхүү хязгаарлалтыг арилгах зорилгоор Caner "
        "ба Hansen (2004) нь хэрэгсэл хувьсагч ба босгыг нэг загварт "
        "нэгтгэсэн IVTR алгоритмыг дэвшүүлэв. Уг арга нь санхүүгийн "
        "хөгжлийн босго, хөгжлийн тусламжийн үр нөлөө, бичил санхүүгийн "
        "өгөөж зэрэг олон салбарт хэрэглэгдсэн боловч хөдөлмөрийн эдийн "
        "засгийн контекстэд өргөн хэрэглэгдээгүй байгаагийн гол жишээ нь "
        "Chakroun (2013)-ын санхүүгийн хөгжил ба тэгш бус байдлын босго "
        "шинжилгээний ажил юм.")
    add_paragraph(doc,
        "Монголын эмпирик уран зохиолд босго загварыг сүүлийн үед хэрэглэж "
        "эхэлсэн бөгөөд Даваажаргал ба Цолмон (2019) нар Монголбанкны "
        "\"Товхимол 14\" судалгааны ажлын цувралд Threshold SVAR загварыг "
        "Монголд анх удаа хэрэглэж, Засгийн газрын өрийн хэмжээний "
        "төсвийн бодлогын үр нөлөөнд нөлөөлж буйг харуулсан. Харин "
        "өрхийн микро өгөгдөл дээр IV ба босго шинжилгээг хослуулсан "
        "бүтээл Монголд одоо хүртэл байхгүй бөгөөд энэхүү судалгаа нь "
        "тэр хоосролыг нөхөхийг зорьж байна.")

    add_heading(doc, "1.4. Монголын хүрээн дэх өмнөх судалгаа", level=2)
    add_paragraph(doc,
        "Монгол Улс дахь боловсролын өгөөжийн эмпирик судалгаа цөөн "
        "тоотой боловч хэд хэдэн чухал бүтээл байна. Pastore (2010) нь "
        "2007–2008 оны өрхийн түүвэрт тулгуурлан залуучуудын боловсролын "
        "өгөөжийг 7–9% орчим хэмээн тогтоосон. Гэвч уг бүтээл нь энгийн "
        "OLS хэлбэрийн үнэлгээгээр хязгаарлагдсан бөгөөд хэрэгсэл "
        "хувьсагч, шугаман бус шинжилгээг хийгээгүй. Түүнчлэн тус "
        "судалгааны цаг үе 2008 оны боловсролын шинэчлэл (10 жилээс 12 "
        "жил болсон), дээд боловсролын хамралт 3 дахин өссөн зэрэг "
        "бүтцийн өөрчлөлтүүдээс өмнөх үеийг хамардаг тул одоогийн "
        "Монгол Улсын нөхцөл байдалтай нийцэхгүй болж байна.")
    add_paragraph(doc,
        "Ай Ар Ай Эм ХХК (2015) нь Боловсрол, Соёл, Шинжлэх Ухаан, Спортын "
        "Яамны захиалгаар \"Хөдөлмөрийн зах зээлийн шинжилгээ\" тайланг "
        "боловсруулж, Ажиллах Хүчний Судалгааны (АХС) 2013 оны өгөгдөлд "
        "тулгуурласан Минсерийн регрессийн үр дүнг танилцуулжээ. Уг "
        "тайлангийн гол дүгнэлт нь дээд боловсролтой иргэд бүрэн дунд "
        "боловсролтой иргэдээс дунджаар 41%, дипломын дараах (магистр, "
        "доктор) зэрэгтэй иргэд 54% илүү цалин авдаг байна гэсэн явдал "
        "юм. Энэхүү тоо Claudio ба Patrinos (2014)-ийн Дэлхийн банкны "
        "судалгааны Монголд тооцоолсон (дунд боловсрол 4.2%, дээд "
        "боловсрол 10.1%)-тай харьцуулах боломжтой. Гэвч Ай Ар Ай Эм-ийн "
        "тайланд IV, босго шинжилгээ хийгдээгүй.")
    add_paragraph(doc,
        "Эндээс үзвэл Монголын эмпирик уран зохиолд боловсролын "
        "өгөөжийг микро өгөгдөл дээр IV, босго хэлбэрээр нэгтгэн "
        "үнэлсэн ажил одоог хүртэл хийгдээгүй байна. Ижил төрлийн "
        "сорилт бусад шилжилтийн эдийн засгуудад ч тохиолдож буй "
        "бөгөөд уг хоосролыг нөхөх нь Монголын хөдөлмөрийн эдийн "
        "засгийн судалгаанд чухал хувь нэмэр оруулах боломжтой.")

    # ─── ХОЁРДУГААР БҮЛЭГ. СУДАЛГААНЫ АРГА ЗҮЙ ───────────────────────────────
    add_heading(doc, "II бүлэг. Судалгааны арга зүй", level=1,
                before_break=True)

    add_paragraph(doc,
        "Энэхүү судалгаанд боловсролын өгөөжийн бодит хэмжээг үнэлэхдээ "
        "дөрвөн шатлалт эконометрик шинжилгээг явуулав. Нэгдүгээр шатанд "
        "Минсерийн сонгодог цалингийн тэгшитгэлийг энгийн OLS аргаар "
        "үнэлж суурь үнэлгээг гаргана. Хоёрдугаар шатанд төрсөн когорт × "
        "аймаг × давалгааны псевдо-панелийн нүднүүд дээр фиксэлсэн ба "
        "санамсаргүй нөлөөний загваруудыг харьцуулан Хаусманы шалгуураар "
        "тохирох загварыг сонгоно. Гуравдугаар шатанд боловсролын "
        "эндогений шинжийг арилгах зорилгоор хоёр шатлалт хамгийн бага "
        "квадратын (2SLS) аргыг хэрэглэнэ. Эцсийн шатанд Caner ба Hansen "
        "(2004)-ийн хэрэгсэл хувьсагчтай босготой регрессийг хэрэгжүүлж, "
        "боловсролын жилийн хэмжээнд бүтцийн хугарал оршиж буй эсэхийг "
        "тогтооно. Ийнхүү шатлан явуулах нь үр дүнгийн тогтвортой "
        "байдлыг баталгаажуулах шаардлагатай болно.")

    add_heading(doc, "2.1. OLS-ийн Минсерийн загвар", level=2)
    add_paragraph(doc,
        "Суурь загвар нь цалингийн логарифмыг боловсролын жил, ажлын "
        "туршлага, туршлагын квадрат болон хяналтын хувьсагчуудаар "
        "тайлбарлах өргөтгөсөн Минсерийн тэгшитгэл юм:")
    if 'EQ_MINCER_EXTENDED' in EQ:
        add_equation(doc, EQ['EQ_MINCER_EXTENDED'], number=2)
    add_paragraph(doc,
        "Энд X нь хяналтын хувьсагчдын багц (хүйс, гэр бүлийн байдал, "
        "суурьшил, өрхийн хэмжээ, эдийн засгийн салбарын ангилал), "
        "μₐ нь аймгийн фиксэлсэн нөлөө, τₜ нь ӨНЭЗС-ийн давалгааны "
        "фиксэлсэн нөлөө юм. Стандарт алдааг аймаг тус бүрээр "
        "кластерлагдсан хэлбэрээр тооцсон бөгөөд энэ нь ижил аймаг "
        "доторх ажиглалтуудын хүрээнд үүсэх боломжит корреляцийг "
        "харгалзан авдаг.")

    add_heading(doc, "2.2. Псевдо-панел FE/RE загвар", level=2)
    add_paragraph(doc,
        "ӨНЭЗС нь давхардаагүй түүвэр судалгаа учраас жинхэнэ "
        "хугацаат панел байгуулах боломжгүй. Иймд Deaton (1985)-ийн "
        "псевдо-панелийн аргаар төрсөн когорт (c, 5-жилийн бүлэглэл) × "
        "аймаг (a) × давалгаа (t) гэсэн гурван хэмжээсээр нүдлэн нэгтгэж, "
        "нийт 692 нүдтэй бүтцийг үүсгэв. Фиксэлсэн нөлөөний (FE) загвар "
        "дараах хэлбэртэй:")
    if 'EQ_FE_PANEL' in EQ:
        add_equation(doc, EQ['EQ_FE_PANEL'], number=3)
    add_paragraph(doc,
        "Энд c тус бүрийн нүдний хувьсагчдын дунджийг дээд зураастайгаар "
        "тэмдэглэсэн болно. Санамсаргүй нөлөөний (RE) загвар нь αca-г "
        "санамсаргүй хүчин зүйл гэж үзнэ. FE ба RE загваруудын "
        "сонголтыг Хаусманы шалгуураар хийв:")
    if 'EQ_HAUSMAN' in EQ:
        add_equation(doc, EQ['EQ_HAUSMAN'], number=4)
    add_paragraph(doc,
        "Хэрэв Хаусманы статистик χ² критик утгаас давсан бол FE "
        "загварыг сонгоно. Энэ нь нүдний тогтмол хүчин зүйлс нь "
        "тайлбарлах хувьсагчтай корреляцитай болохыг илтгэдэг.")

    add_heading(doc, "2.3. Хоёр шатлалт хамгийн бага квадратын (2SLS) арга", level=2)
    add_paragraph(doc,
        "Боловсрол нь хувь хүний ажиглагдахгүй ур чадвар, авъяас, гэр "
        "бүлийн нөхцөлтэй нягт холбоотой тул OLS үнэлгээ хазайлттай "
        "болох магадлалтай. Энэхүү эндогений шинжийг арилгахын тулд "
        "боловсролын сонголтод нөлөөлдөг боловч цалинд зөвхөн "
        "боловсролоор дамжин нөлөөлдөг гуравдагч хүчин зүйлсийг "
        "хэрэгсэл хувьсагч болгон хэрэглэх шаардлагатай.")
    add_paragraph(doc,
        "Энэхүү судалгаанд гурван хэрэгсэл хувьсагчийг сонгов. Эхнийх "
        "нь төрсөн аймгийн 21 ангилал хувьсагч юм. Аймгийн фиксэлсэн "
        "нөлөөг хянасны дараа сумын дотоод ялгаа нь хувь хүний авъяас "
        "чадвараас хамаарахгүй, харин тухайн сумын боловсролын "
        "хүртээмжтэй холбоотой байдаг. Хоёрдох хэрэгсэл нь төрсөн "
        "сумын төвөөс Улаанбаатар хүртэлх зайн логарифм — хотоос хол "
        "сумд сурах боломж бага, тээврийн зардал өндөр байсан нь "
        "боловсролын сонголтыг шууд хязгаарладаг. Гуравдах хэрэгсэл нь "
        "тухайн хүн 12 настай байх жилд төрсөн суманд ногдож байсан "
        "ЕБС-ийн багш/сурагчийн харьцаа бөгөөд 1212.mn нээлттэй "
        "мэдээллийн сангаас татав. Duflo (2001)-ийн Индонезийн "
        "судалгаатай ижил логикийг Монгол Улсын нөхцөлд шилжүүлэн "
        "хэрэглэж буй явдал юм.")
    add_paragraph(doc,
        "2SLS үнэлгээ нь хоёр шатаас бүрдэнэ. Эхний шатанд эндогений "
        "регрессорыг хэрэгсэл хувьсагчаар багасгасан хэлбэрт (reduced-"
        "form) шилжүүлнэ:")
    if 'EQ_IV_STAGE1' in EQ:
        add_equation(doc, EQ['EQ_IV_STAGE1'], number=5)
    add_paragraph(doc,
        "Хоёр дахь шатанд эхний шатны тохирсон утгыг үндсэн тэгшитгэлд "
        "орлуулж цалингийн логарифмыг үнэлнэ:")
    if 'EQ_IV_STAGE2' in EQ:
        add_equation(doc, EQ['EQ_IV_STAGE2'], number=6)
    add_paragraph(doc,
        "Хэрэгсэл хувьсагчийн хүчийг Клейберген–Паапын F статистикаар "
        "(Staiger-Stock шалгуураар 10-аас их байх ёстой) шалгана. "
        "Хэрэгсэл хувьсагчийн тоо эндогений регрессорын тооноос илүү "
        "тохиолдолд Хансены J-тестээр хэт тодорхойлогдсон байдлыг "
        "шалгаж, Андерсон–Рубины тэсвэртэй итгэлийн интервалыг сул "
        "хэрэгсэл хувьсагчийн тохиолдолд нэмэлтээр тооцдог.")

    add_heading(doc, "2.4. Хэрэгсэл хувьсагчтай босготой регресс (Caner–Hansen IVTR)", level=2)
    add_paragraph(doc,
        "Судалгааны гол аргачлал нь Caner ба Hansen (2004)-ийн "
        "хэрэгсэл хувьсагчтай босготой регресс юм. Боловсролын өгөөж "
        "нь хувь хүний боловсролын түвшнээс шугаман бус байдлаар "
        "хамаарах магадлалтай бөгөөд энэ тохиолдолд өгөгдлийг тодорхой "
        "босго утга γ-ын доод ба дээд талд хоёр регимд хуваан, регим "
        "тус бүрийн өгөөжийн коэффициент β-г тусад нь үнэлэх нь "
        "бодитой юм:")
    if 'EQ_IVTR' in EQ:
        add_equation(doc, EQ['EQ_IVTR'], number=7)
    add_paragraph(doc,
        "Энд q = educ нь босго хувьсагч бөгөөд γ нь тодорхойлогдох "
        "босгоны утга юм. educ нь эндогений регрессор тул Z хэрэгсэл "
        "хувьсагчаар багасгагдана.")
    add_paragraph(doc,
        "Үнэлгээний алгоритм дараах алхмуудаар хэрэгжинэ. Эхлээд "
        "эндогений регрессорыг хэрэгсэл хувьсагчаар багасгасан хэлбэрт "
        "шилжүүлнэ. Дараа нь боломжит γ-ийн утгуудыг түүврийн 15–85 "
        "хувийн хооронд сүлжээн хайлтаар (grid search) шалгаж, регим "
        "тус бүрт 2SLS үнэлгээ хийж, концентрацилагдсан алдааны "
        "квадратуудын нийлбэрийг (SSR) тооцно. Хамгийн бага SSR-тэй "
        "γ-г оновчтой босго γ* болгон сонгоно. Босгоны статистик "
        "ач холбогдлыг SupWald шалгуураар тогтооно:")
    if 'EQ_SUPWALD' in EQ:
        add_equation(doc, EQ['EQ_SUPWALD'], number=8)
    add_paragraph(doc,
        "Энд SSR₀ нь босгогүй (нэг регимтэй) загварын SSR, SSR(γ) нь "
        "тодорхой γ-д босгосон хоёр регимт загварын SSR юм. SupWald "
        "статистикийн хуваарилалт стандарт бус тул Радемахерийн жинтэй "
        "дахин түүвэрлэх шалгуур (wild bootstrap)-ыг 1,000 удаагийн "
        "давталтаар явуулж p-утгыг тооцов.")
    add_paragraph(doc,
        "Caner–Hansen-ийн алгоритмыг хэрэгжүүлэх R эсвэл Python багц "
        "одоогоор CRAN болон PyPI дээр байхгүй байгаа нь практик "
        "сорилт болов. Иймд уг алгоритмыг R хэл дээр 180 орчим мөртэй "
        "гарын авлагын код болгон бичиж, параметр нь мэдэгдэж буй "
        "симуляцитай өгөгдөл дээр нэгж шалгалт (unit test) явуулан "
        "баталгаажуулав. Энэ нь дараагийн судлаачдад дахин хэрэглэгдэх "
        "код суурийг бий болгосон давхар ашиг тустай ажил болов.")

    add_heading(doc, "2.5. Тогтвортой байдлын шинжилгээний бүтэц", level=2)
    add_paragraph(doc,
        "Үндсэн үр дүнгийн тогтвортой байдлыг зургаан төрлийн "
        "шинжилгээгээр баталгаажуулсан болно. Нэгдүгээрт, альтернатив "
        "хяналтын хувьсагчдыг нэмж регрессийг дахин үнэлэв. Хоёрдугаарт, "
        "түүврийг дэд бүлгүүдэд (эрэгтэй, эмэгтэй, хот, хөдөө, залуу, "
        "ахмад насны ажилтнууд) хуваан регрессийг тус бүрээр хийж "
        "коэффициентийн тогтвортой байдлыг шалгасан. Гуравдугаарт, өөр "
        "хэрэгсэл хувьсагчдыг ашиглан 2SLS-ын коэффициент үл "
        "хувирахыг баталгаажуулсан. Дөрөвдүгээрт, хяналтын тест "
        "(placebo) хийж — ЕБС-ийн суурь боловсрол дээр IV-ийн нөлөө "
        "мэдэгдэхүйц байх ёсгүй гэсэн таамаглалыг шалгасан. "
        "Тавдугаарт, цалинтай ажилтан болох сонголтын хазайлтыг "
        "арилгах үүднээс Хекманы сонголтын засварыг явуулсан. "
        "Зургаадугаарт, альтернатив босго утгууд (12, 14 жил) дээр "
        "регим хоорондын ялгааны тогтвортой байдлыг шалгасан.")

    # ─── ГУРАВДУГААР БҮЛЭГ. ӨГӨГДӨЛ ──────────────────────────────────────────
    add_heading(doc, "III бүлэг. Өгөгдөл, хувьсагчдын бүтэц",
                level=1, before_break=True)

    add_heading(doc, "3.1. Үндсэн эх сурвалж: ӨНЭЗС", level=2)
    add_paragraph(doc,
        "Энэхүү судалгааны үндсэн өгөгдлийн сурвалж нь Үндэсний "
        "Статистикийн Хороо (ҮСХ)-ноос тогтмол явуулдаг Өрхийн нийгэм, "
        "эдийн засгийн судалгаа (ӨНЭЗС) юм. ӨНЭЗС нь Монгол Улсын "
        "өрхийн орлого, зарлага, хөдөлмөр эрхлэлт, боловсрол, хүн амын "
        "бүтэц зэрэг нийгэм эдийн засгийн үзүүлэлтүүдийг хамарсан "
        "улсын хэмжээний төлөөллөө харгалзсан түүвэр судалгаа бөгөөд "
        "Дэлхийн банк, Олон улсын валютын сан, Нэгдсэн Үндэстний "
        "Байгууллагын статистикийн агентлагуудаар баталгаажуулсан арга "
        "зүйгээр явагддаг.")
    add_paragraph(doc,
        "Энэхүү судалгаанд 2016, 2018, 2020, 2021, 2024 оны таван "
        "давалгаа ашиглагдсан бөгөөд нийт 272,096 хувь хүний "
        "бүртгэлтэйгээр шинжилгээнд оруулав. 2016 ба 2018 оны "
        "давалгаануудад төрсөн аймгийн мэдээлэл байхгүй байсан тул "
        "тэдгээрийг зөвхөн робастнес шинжилгээний дэд түүвэрт "
        "хамруулав. Харин 2020, 2021, 2024 оны гурван давалгаа нь "
        "төрсөн аймгийн хувьсагчтай тул IV шинжилгээний үндсэн түүвэр "
        "болов. Эцсийн шинжилгээнд 25–60 насны цалинтай ажилтан 49,366 "
        "хүнийг хамруулав.")

    add_heading(doc, "3.2. Нэмэлт эх сурвалж: 1212.mn нээлттэй мэдээллийн сан", level=2)
    add_paragraph(doc,
        "Судалгааны гуравдах хэрэгсэл хувьсагчийг бүтээхдээ ҮСХ-ны "
        "1212.mn нээлттэй мэдээллийн сангийн PX-Web JSON API-аар "
        "ерөнхий боловсролын сургуулийн статистикийн хүснэгтүүдийг "
        "татав. 21 аймаг, 330 сумын 2000–2024 оны 25 жилийн хугацаанд "
        "ЕБС-ийн багш, сурагчийн тоог цуглуулж, багш/сурагчийн "
        "харьцааг дараах хэлбэрээр тооцов:")
    if 'EQ_RATIO' in EQ:
        add_equation(doc, EQ['EQ_RATIO'], number=9)
    add_paragraph(doc,
        "Нэмэлт хэрэгсэл хувьсагч болох Улаанбаатар хүртэлх зайг сумын "
        "төвөөс нийслэл хүртэлх шулуун зайгаар, ҮСХ-ны газарзүйн "
        "мэдээллийн сан ба Google Maps API-аас тооцож тогтоов. Зайн "
        "логарифмыг ашигласан нь гадна утгын нөлөөг багасгах зорилготой.")

    add_heading(doc, "3.3. Өгөгдлийн боловсруулалт, цэвэрлэгээ", level=2)
    add_paragraph(doc,
        "Таван давалгааг нэгтгэхэд хувьсагчуудын нэршил жигд бус байсан "
        "нь боловсруулалтын гол сорилт байв. Цалингийн хувьсагч ӨНЭЗС-"
        "ийн эхний хоёр давалгаанд (2016, 2018) нэг код, сүүлийн гурван "
        "давалгаанд (2020, 2021, 2024) өөр код дор оршиж байсан тул "
        "нэгдсэн wage_monthly хувьсагчид тохирох хөрвүүлэлт хийв. "
        "Боловсролын жилийн талаар сүүлийн гурван давалгаанд жилийн "
        "тоо шууд байсан бол 2016, 2018 оны давалгаанд зөвхөн боловсролын "
        "түвшний ангилал байсан учраас боловсролын 6 түвшнийг "
        "эквивалент жилд буулгах хөрвүүлэлтийг хийж нийт 93,420 "
        "ажиглалтыг нөхөв.")
    add_paragraph(doc,
        "Нөхсөн өгөгдлөөс үүдэх сонголтын хазайлтын эрсдэлийг арилгах "
        "зорилгоор Хекманы хоёр шатлалт аргаар шалгахад Милсийн "
        "урвуу харьцаа λ = -0.041 (p > 0.5) буюу статистикийн хувьд "
        "ач холбогдолгүй гарсан. Энэ нь нөхөлт хийсэн ажиглалтууд "
        "нөхөлтгүй ажиглалтуудтай системтэй ялгаагүй бөгөөд "
        "шинжилгээний түүвэрт хазайлт үүсээгүй болохыг харуулж байна. "
        "Үндсэн түүвэрт 25–60 насны цалинтай ажилтныг оруулж, сарын "
        "цалингийн 84,000 төгрөгөөс доош болон 5,000,000 төгрөгөөс "
        "дээш утгуудыг гаднаас нь онцгой гараасан (outlier) тул "
        "түүврээс хасав.")

    add_heading(doc, "3.4. Описатив статистик", level=2)
    add_paragraph(doc,
        "Шинжилгээний түүврийн үндсэн үзүүлэлтүүдийг Хүснэгт 1-д "
        "танилцуулав. Боловсролын дундаж жил 11.89 буюу медиан 12 "
        "жил байгаа нь Монголын хөдөлмөр эрхэлж буй насанд хүрэгчдийн "
        "ихэнх нь бүрэн дунд боловсролтой байгааг харуулж байна. "
        "Цалингийн хуваарилалт log-нормал тархалттай ойролцоо бөгөөд "
        "дундаж 861 мянган төгрөг, медиан 700 мянган төгрөг гэсэн "
        "утгатай байна.")

    # Table 1.1
    add_table_caption(doc, "1", "Шинжилгээний түүврийн үндсэн үзүүлэлт")
    build_table(doc,
        headers=["Хувьсагч", "N", "Дундаж", "Ст. хазайлт", "Медиан", "Min", "Max"],
        rows=[
            ("Log (сарын цалин)", "49,366", "13.47", "0.62", "13.46", "11.34", "15.43"),
            ("Сарын цалин (₮)", "49,366", "861,438", "606,250", "700,000", "84,000", "5,000,000"),
            ("Боловсролын жил", "49,357", "11.89", "2.89", "12", "0", "22"),
            ("Нас (жил)", "49,366", "39.44", "9.09", "39", "25", "60"),
            ("Туршлага (жил)", "49,366", "21.54", "10.02", "21", "0", "54"),
            ("Хүйс (эмэгтэй=1)", "49,366", "0.50", "0.50", "0", "0", "1"),
            ("Хот (УБ=1)", "49,366", "0.32", "0.47", "0", "0", "1"),
            ("Өрхийн хэмжээ", "49,366", "4.03", "1.56", "4", "1", "16"),
        ],
        col_widths_cm=[4.0, 1.5, 1.5, 2.0, 1.8, 1.5, 1.7],
        align='right')
    add_table_source(doc, "Эх сурвалж: Оюутны тооцоолол")

    add_paragraph(doc,
        "ӨНЭЗС-ийн давалгаа тус бүрээр тооцсон үндсэн үзүүлэлтүүдийг "
        "Хүснэгт 2-т танилцуулав. Хүснэгтээс харахад цалингийн нэрлэсэн "
        "хэмжээ 2016 оны 544 мянган төгрөгөөс 2024 оны 1,526 мянган "
        "төгрөг хүртэл 2.8 дахин өсчээ. Энэ нь Монгол Улсын "
        "хөдөлмөрийн зах зээлд нэрлэсэн цалингийн өсөлт огцом явагдсаныг "
        "илэрхийлдэг. Түүнчлэн эмэгтэй ажилтны эзлэх хувь 49.2–50.7%-"
        "ийн хооронд тогтворжсон бол хотын ажилтны эзлэх хувь 64.6–"
        "73.9%-ийн хооронд хэлбэлзэж байна.")

    # Table 1.2
    add_table_caption(doc, "2",
                      "Түүврийн үндсэн үзүүлэлт ӨНЭЗС-ийн давалгаа тус бүрээр")
    build_table(doc,
        headers=["Давалгаа", "N", "Дундаж боловсрол (жил)",
                 "Дундаж цалин (₮)", "Медиан цалин (₮)",
                 "Эмэгтэй (%)", "Хотод (%)"],
        rows=[
            ("2016", "10,428", "11.8", "543,714", "500,000", "50.7", "67.1"),
            ("2018", "10,955", "11.7", "610,882", "520,000", "49.2", "68.2"),
            ("2020", "11,045", "11.9", "795,270", "700,000", "49.7", "68.0"),
            ("2021", "6,713", "12.2", "859,801", "750,000", "49.4", "73.9"),
            ("2024", "10,225", "12.0", "1,526,462", "1,400,000", "49.8", "64.6"),
        ],
        col_widths_cm=[1.8, 1.8, 2.5, 2.5, 2.5, 1.7, 1.7],
        align='right')
    add_table_source(doc, "Эх сурвалж: Оюутны тооцоолол")

    add_paragraph(doc,
        "Боловсролын жилийн тархалтыг ӨНЭЗС-ийн давалгаа тус бүрээр "
        "Зураг 1-т харуулав. Зургаас харахад тархалт нь бүх давалгаанд "
        "хоёр оргилтой (bimodal) онцлогтой байна: нэгэн оргил 10–12 "
        "жилийн боловсролтой (бүрэн дунд төгссөн) бүлэгт, нөгөөх нь "
        "14–16 жилийн боловсролтой (коллеж, бакалавр) бүлэгт оршдог. "
        "Энэхүү бүтэц нь Монголын боловсролын тогтолцооны түүхэн "
        "өөрчлөлттэй нягт холбоотой: 1995 он хүртэл 10 жилийн ерөнхий "
        "боловсролын систем, 1995–2008 онд 11 жилийн шинэчилсэн "
        "систем, 2008 оноос хойш 12 жилийн систем үйлчилж ирсэн. "
        "Иймд одоогийн 40-өөс дээш насны ажилтнуудын дунд 10–11 "
        "жилийн боловсрол зонхилдог бол 35-аас доош насны залуу "
        "ажилтнууд 12 жилийн бүрэн дунд боловсролтой байна.")
    add_paragraph(doc,
        "Хоёр оргилын хооронд буюу 13 жилийн боловсролтой иргэдийн "
        "тоо харьцангуй цөөн байгаа нь нүдэнд тодорхой харагдаж байна. "
        "Энэ хоосрол нь бүрэн дунд боловсролоос дээд боловсрол руу "
        "шилжих замд институцийн саад — тухайлбал их, дээд сургуулийн "
        "элсэлтийн шалгалт, сургалтын төлбөрийн саад, нийгэм эдийн "
        "засгийн хязгаарлалт — байгааг илэрхийлж байна. Үүнийг "
        "Мөнхнасан ба Нарандалай (2013)-ийн Монголын хүн амын орлогын "
        "шинжилгээнд мөн дурдсан байдаг. Иймд бүрэн дунд ба дээд "
        "боловсролын хоорондын шилжилт нь боловсролын өгөөжийн "
        "бүтцийн хугарлын нэг магадлалт цэг болох урьдчилсан "
        "нотолгоо юм.")

    # Figure 2.1
    add_image_centered(doc, FIGURES / "f1_education_distribution.png", width_cm=15)
    add_figure_caption(doc, "1",
                       "Боловсролын жилийн тархалт ӨНЭЗС-ийн давалгаа бүрээр")
    add_figure_source(doc, "Эх сурвалж: Оюутны тооцоолол")

    add_paragraph(doc,
        "Боловсролын жил ба цалингийн логарифмын хоорондох харилцааг "
        "2024 оны давалгаан дээр Зураг 2-т харуулав. Шугаман OLS "
        "шугам нь нийт түүвэрт ойролцоогоор 0.06-ын налуутай байхад "
        "бодит цэгийн тархалт нь 13 жилийн орчимд бүтцийн хугарал "
        "үүсгэж буй нь нүдэнд харагдаж байна. Энэ нь босго шинжилгээний "
        "таамаглалыг нүдний зүйн нотолгоогоор баталгаажуулах анхны сэжим болж, "
        "дараагийн IVTR үнэлгээний суурь нотолгоо болов.")

    # Figure 2.2
    add_image_centered(doc, FIGURES / "f2_wage_education_scatter.png", width_cm=14)
    add_figure_caption(doc, "2",
                       "Боловсролын жил ба логаритмчилсан цалингийн хамаарал")
    add_figure_source(doc, "Эх сурвалж: Оюутны тооцоолол")

    # ─── ДӨРӨВДҮГЭЭР БҮЛЭГ. ҮР ДҮН ───────────────────────────────────────────
    add_heading(doc, "IV бүлэг. Шинжилгээний үр дүн",
                level=1, before_break=True)

    add_heading(doc, "4.1. OLS ба фиксэлсэн нөлөөний үнэлгээ", level=2)
    add_paragraph(doc,
        "Суурь OLS үнэлгээний үр дүнг Хүснэгт 3-т танилцуулав. Энгийн "
        "OLS загварт боловсролын нэг жилийн өгөөж нь β = 0.0736 буюу "
        "7.6% гэж тогтоогдсон. Энэ нь хувь хүний боловсролын нэмэлт нэг "
        "жил цалинг 7.6%-иар өсгөдөг гэсэн үг юм. Хяналтын хувьсагчуудыг "
        "нэмэхэд β = 0.0791 (8.2%) хүртэл өссөн бол аймаг болон "
        "давалгааны фиксэлсэн нөлөөг оруулахад β = 0.0661 (6.8%) хүртэл "
        "буурсан. Иймд аймгийн газарзүйн болон үеийн тогтмол нөлөө нь "
        "боловсрол-цалингийн хамаарлыг хиймэлээр дээш чиглэсэн хазайлтад "
        "орох хандлагатай байгаа нь харагдаж байна. Тодорхойлогдох "
        "коэффициент R² нь 0.111-ээс 0.526 хүртэл огцом өсөж, аймгийн "
        "тогтмол нөлөө нь цалингийн хэлбэлзлийн дийлэнх хэсгийг "
        "тайлбарлаж чадсан нь харагдаж байна.")

    # Table 3.1
    add_table_caption(doc, "3",
                      "Энгийн OLS ба фиксэлсэн нөлөөтэй үнэлгээний үр дүн")
    build_table(doc,
        headers=["Үзүүлэлт", "(1) OLS энгийн", "(2) OLS + хяналт",
                 "(3) OLS + FE"],
        rows=[
            ("Боловсролын жил (β)", "0.0736***", "0.0791***", "0.0661***"),
            ("", "(0.0011)", "(0.0010)", "(0.0019)"),
            ("Туршлага", "0.0276***", "0.0307***", "0.0308***"),
            ("Туршлагын квадрат", "−0.00059***", "−0.00065***", "−0.00071***"),
            ("Аймгийн FE", "Үгүй", "Үгүй", "Тийм"),
            ("Давалгааны FE", "Үгүй", "Үгүй", "Тийм"),
            ("Бусад хяналт", "Үгүй", "Тийм", "Тийм"),
            ("N", "49,357", "49,357", "49,357"),
            ("R²", "0.111", "0.172", "0.526"),
        ],
        col_widths_cm=[5.0, 3.3, 3.3, 3.3],
        align='center')
    # Note row
    p_note = doc.add_paragraph()
    p_note.paragraph_format.space_before = Pt(3)
    p_note.paragraph_format.space_after = Pt(0)
    p_note.paragraph_format.line_spacing = 1.1
    p_note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p_note.add_run("Тэмдэглэл: ")
    set_run_props(r, size=9, italic=True, bold=True)
    r = p_note.add_run(
        "Хаалтан доторх тоо нь аймгаар кластерлесан стандарт алдаа. "
        "Статистик ач холбогдол: *p<0.10, **p<0.05, ***p<0.01.")
    set_run_props(r, size=9, italic=True)
    add_table_source(doc, "Эх сурвалж: Оюутны тооцоолол")

    add_paragraph(doc,
        "Псевдо-панелийн фиксэлсэн нөлөөний үнэлгээгээр β_FE = 0.1627 "
        "(стандарт алдаа 0.0124), санамсаргүй нөлөөний үнэлгээгээр "
        "β_RE = 0.2028 (0.0278) гарав. Хаусманы шалгуурын χ² = 15,282, "
        "p < 0.001 буюу хоёр загварын үнэлгээний параметрүүд статистикийн "
        "хувьд ялгаатай тул фиксэлсэн нөлөөний загварыг сонгосон. "
        "Псевдо-панелийн түвшинд өгөөжийн үнэлгээ хувь хүний түвшний "
        "үнэлгээнээс мэдэгдэхүйц өндөр гарсан нь нэгдмэлжсэн когортын "
        "түвшинд хүмүүн капиталын хуримтлалын нийт нөлөө илүү тод "
        "илэрдэгтэй холбоотой (Deaton, 1985).")

    add_heading(doc, "4.2. Хоёр шатлалт (2SLS) үнэлгээ", level=2)
    add_paragraph(doc,
        "OLS үнэлгээний эндогений хазайлтыг арилгах зорилгоор гурван "
        "өөр хэрэгсэл хувьсагчтайгаар 2SLS үнэлгээг явуулж, үр дүнг "
        "Хүснэгт 4-т нэгтгэн танилцуулав. Хэрэгсэл хувьсагч тус бүр "
        "өөр өөр түүвэрт хамаарч, өөр өөр эхний шатны хүчтэй байх тул "
        "харьцуулах боломжийг олгож байна.")

    # Table 3.2
    add_table_caption(doc, "4", "Хэрэгсэл хувьсагчийн (IV) үнэлгээний үр дүн")
    build_table(doc,
        headers=["Үзүүлэлт", "(1) OLS",
                 "(2) 2SLS төрсөн аймаг", "(3) 2SLS ln(dist)",
                 "(4) 2SLS ЕБС", "(5) Overident."],
        rows=[
            ("Боловсролын жил (β)", "0.0574***", "0.1133***", "0.2055***",
             "0.0944***", "0.1043***"),
            ("", "(0.0025)", "(0.0068)", "(0.0389)", "(0.0308)", "(0.0054)"),
            ("Эхний шатны F", "—", "10.46", "17.77", "110.02", "—"),
            ("Hansen J (p-утга)", "—", "—", "—", "—", "0.035"),
            ("N", "12,020", "12,020", "7,746", "1,367", "7,746"),
        ],
        col_widths_cm=[3.8, 2.3, 2.6, 2.4, 2.3, 2.5],
        align='center')
    p_note = doc.add_paragraph()
    p_note.paragraph_format.space_before = Pt(3)
    p_note.paragraph_format.space_after = Pt(0)
    p_note.paragraph_format.line_spacing = 1.1
    p_note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p_note.add_run("Тэмдэглэл: ")
    set_run_props(r, size=9, italic=True, bold=True)
    r = p_note.add_run(
        "Хаалтан доторх тоо нь аймгаар кластерлесан стандарт алдаа. "
        "Үндсэн түүвэр — ӨНЭЗС 2020+2021+2024. ***p<0.01, **p<0.05, "
        "*p<0.10.")
    set_run_props(r, size=9, italic=True)
    add_table_source(doc, "Эх сурвалж: Оюутны тооцоолол")

    add_paragraph(doc,
        "Хүснэгт 4-өөс харвал хэрэгсэл хувьсагчийн 2SLS үнэлгээ нь OLS "
        "үнэлгээнээс мэдэгдэхүйц өндөр гарсан нь гол олдвор юм. "
        "Боловсролын нэг жилийн өгөөж OLS-оор 5.7% байсан бол төрсөн "
        "аймгийн IV-тэй 2SLS-аар 11.3% хүртэл өсчээ. Энэ нь OLS "
        "үнэлгээний β нь бодит өгөөжийн ойролцоогоор хагасыг л "
        "илэрхийлж байсан буюу 49%-ийн хазайлт үүссэн гэсэн үг юм. "
        "Өөрөөр хэлбэл, OLS-ын доош чиглэсэн хэмжилтийн хазайлт "
        "(attenuation bias) нь ур чадварын хомсдлоос үүдэлтэй дээш "
        "чиглэсэн хазайлтаас давамгайлж байгааг харуулж байна. "
        "Энэхүү үр дүн нь Card (1993)-ын Хойд Америкт хийсэн сонгодог "
        "судалгаанд тогтоогдсон 25–60%-ийн хазайлттай ойролцоо байгаа "
        "нь онцлох нотолгоо болов.")
    add_paragraph(doc,
        "Хэрэгсэл хувьсагч тус бүрийн эхний шатны F статистикийг "
        "харахад Staiger–Stock-ын шалгуурыг (F > 10) бүгд давсан "
        "байна. ЕБС-ийн багш/сурагчийн харьцааны хэрэгсэл хувьсагч нь "
        "хамгийн хүчтэй шалтгаалалттай (F = 110.02) бол төрсөн аймгийн "
        "хэрэгсэл хувьсагч шалгуурын босго дээр (F = 10.46) байгаа "
        "нь Андерсон–Рубины сул хэрэгсэлд тэсвэртэй итгэлийн "
        "интервалыг нэмэлтээр тайлагнах үндэслэл болов. Ву–Хаусманы "
        "эндогений тестийн F = 22.1 (p < 0.001) нь OLS ба 2SLS "
        "коэффициентүүд статистикийн хувьд ялгаатай болохыг баталж, "
        "IV аргыг хэрэглэх шаардлагатай болохыг эмпирикээр нотолж "
        "байна.")

    add_heading(doc, "4.3. Хэрэгсэл хувьсагчтай босготой регресс — үндсэн үр дүн",
                level=2)
    add_paragraph(doc,
        "Энэхүү судалгааны хамгийн чухал үр дүн нь Caner ба Hansen "
        "(2004)-ийн IVTR аргыг Монгол Улсын өрхийн микро өгөгдөлд анх "
        "удаа хэрэглэсэн үр дүн юм. Сүлжээн хайлтын (grid search) "
        "явцад боловсролын 10–15 жилийн γ хувьсагчийн утгууд дээр "
        "концентрацилагдсан алдааны квадратын нийлбэрийг тооцож, "
        "хамгийн бага утгыг γ* = 13 жил дээр олсон болно. Уг цэг "
        "Монгол Улсын боловсролын тогтолцоонд тодорхой утгатай: 12-р "
        "анги төгсгөсний дараа мэргэжлийн сургалт, МСҮТ, коллежийн "
        "эхний жил, их сургуулийн бэлтгэл курс зэрэг сурах нэмэлт "
        "жил юм.")

    # Figure 3.1
    add_image_centered(doc, FIGURES / "f3_threshold_profile.png", width_cm=13)
    add_figure_caption(doc, "3",
        "Босгоны сүлжээн хайлтын төвлөрсөн алдааны квадратын нийлбэр")
    add_figure_source(doc, "Эх сурвалж: Оюутны тооцоолол")

    # Table 3.3
    add_table_caption(doc, "5",
        "Хэрэгсэл хувьсагчтай босготой регресс (IVTR) — үндсэн үр дүн")
    build_table(doc,
        headers=["Үзүүлэлт", "Утга"],
        rows=[
            ("Оновчтой босго γ*", "13 жил"),
            ("Доод регимийн налуу β₁ (educ ≤ 13)", "0.0533"),
            ("β₁-ийн стандарт алдаа", "(0.0272)"),
            ("Нэг жилийн өгөөж (регим 1)", "5.5%"),
            ("Дээд регимийн налуу β₂ (educ > 13)", "0.1650"),
            ("β₂-ийн стандарт алдаа", "(0.0307)"),
            ("Нэг жилийн өгөөж (регим 2)", "17.9%"),
            ("SupWald статистик", "152.20"),
            ("Wild bootstrap p-утга (1,000 давталт)", "< 0.001"),
            ("Доод регимийн N", "6,144"),
            ("Дээд регимийн N", "5,876"),
            ("Нийт N", "12,020"),
        ],
        col_widths_cm=[10.0, 6.0], align='left', first_col_left=True)
    p_note = doc.add_paragraph()
    p_note.paragraph_format.space_before = Pt(3)
    p_note.paragraph_format.space_after = Pt(0)
    p_note.paragraph_format.line_spacing = 1.1
    p_note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p_note.add_run("Тэмдэглэл: ")
    set_run_props(r, size=9, italic=True, bold=True)
    r = p_note.add_run(
        "Wild bootstrap нь Радемахерийн жингүүдтэй 1,000 борлуулалтаар "
        "үүсгэгдсэн. Үндсэн түүвэр ӨНЭЗС 2020+2021+2024.")
    set_run_props(r, size=9, italic=True)
    add_table_source(doc, "Эх сурвалж: Оюутны тооцоолол")

    add_paragraph(doc,
        "Регим тус бүрийн өгөөжийн зөрүүг Зураг 4-т харьцуулан үзүүлэв.")
    add_image_centered(doc, FIGURES / "f4_regime_slopes.png", width_cm=12)
    add_figure_caption(doc, "4",
        "Босгоор тусгаарлагдсан боловсролын IV өгөөж (Caner–Hansen IVTR)")
    add_figure_source(doc, "Эх сурвалж: Оюутны тооцоолол")

    add_paragraph(doc,
        "Судалгааны хамгийн чухал эмпирик олдвор нь 13 жилийн босгон "
        "дээр боловсролын өгөөжийн огцом үсрэлт юм. Босгоноос доош "
        "(educ ≤ 13) нэг жилийн өгөөж ердөө 5.5%, харин босгоноос дээш "
        "(educ > 13) 17.9% болж 3.3 дахин өсч байна. Өөрөөр хэлбэл, "
        "боловсролын нэмэлт нэг жил цалингийн түвшинд үзүүлэх нөлөө нь "
        "босгоны хоёр талд 12.4 нэгж хувиар ялгаатай байна. Энэ нь "
        "Монгол Улсын хөдөлмөрийн зах зээлд Jaeger ба Page (1996)-ийн "
        "\"дипломын нэмэгдэл\" үзэгдэл (sheepskin effect) эмпирикээр "
        "оршин буйг анх удаа нотолж буй чухал баримт юм.")
    add_paragraph(doc,
        "Бүтцийн хугарлын статистик ач холбогдлыг SupWald шалгуураар "
        "шалгахад утга нь 152.20 болж, стандарт хи-квадратын критик "
        "утгуудаас олон дахин давсан. Түүнчлэн 1,000 удаагийн "
        "Радемахерийн жинтэй дахин түүвэрлэх шалгуурын (wild bootstrap) "
        "p-утга 0.001-ээс бага байгаа нь босго тохиолдлын шинжгүй "
        "статистикийн хувьд маш өндөр ач холбогдолтой болохыг "
        "баталгаажуулав. Доод регимд 6,144 хүн, дээд регимд 5,876 "
        "хүн хуваарилагдав.")
    add_paragraph(doc,
        "Доод регимд (educ ≤ 13) орсон ажилтнууд нь ЕБС төгсөгчид "
        "болон мэргэжлийн богино хугацааны бэлтгэл бүхий хүмүүс юм. "
        "Тэдний 5.5%-ийн нэг жилийн өгөөж нь хөдөлмөрийн зах зээлийн "
        "ерөнхий үйлдвэрлэлийн бүтээмжийн хэвийн үнэлгээг илэрхийлж "
        "байна. Энэхүү харьцангуй бага өгөөж нь Монголын ЕБС-ийн "
        "дунд шатны боловсрол нь одоогийн хөдөлмөрийн зах зээлд "
        "шаардагдаж буй ур чадваруудтай (мэдээллийн технологи, "
        "санхүү, гадаад хэл) бүрэн нийцэж чадахгүй болох магадлалтайг "
        "илтгэж байна.")
    add_paragraph(doc,
        "Харин дээд регимд (educ > 13) орсон ажилтнууд нь коллеж "
        "болон бакалаврын дипломтой хүмүүс юм. Тэдэнд нэмэлт нэг "
        "жилийн боловсролын өгөөж 17.9% болж, энэ нь дэлхийн дундаж "
        "(Psacharopoulos ба Patrinos, 2018-аар 14.6%) түвшнийг даван "
        "гарч байна. Уг өндөр өгөөж нь Spence (1973)-ийн дохионы "
        "онолтой нийцэж буй бөгөөд дипломыг ажил олгогч нар "
        "ажилтны бүтээмжийн дохио болгон ашигладаг болохыг харуулж "
        "байна. Монгол Улсад төрийн алба, эрүүл мэнд, боловсрол, "
        "хууль эрх зүй, санхүүгийн салбарт бакалавр түүнээс дээш "
        "дипломыг хатуу шалгуур болгодог учир дээд регимд илэрч буй "
        "дипломын нэмэгдэл нь зөвхөн ур чадварын хуримтлалын нөлөө "
        "бус, институцийн шаардлагаас үүдэлтэй ач холбогдолтой юм.")

    add_paragraph(doc,
        "OLS ба IV үнэлгээний давалгаа хоорондын харьцуулалтыг "
        "Зураг 5-т танилцуулав. OLS үнэлгээ нь давалгаа тус бүрд "
        "4.9–8.9%-ийн хооронд хэлбэлзэж аажим буурах хандлагатай "
        "байгаа нь сүүлийн жилүүдэд шинэ ажилтнуудын боловсрол-"
        "цалингийн шугаман хамаарал ядавхаж байгааг илтгэж байна. "
        "Харин IV үнэлгээ нь ойролцоогоор 11% түвшинд тогтвортой "
        "хадгалагдаж буй нь эндогений хазайлтыг засаж үнэлсэн "
        "бодит өгөөж нь цаг хугацааны хувьд тогтвортой болохыг "
        "харуулна.")

    add_image_centered(doc, FIGURES / "f5_ols_vs_iv.png", width_cm=14)
    add_figure_caption(doc, "5",
        "OLS ба IV үнэлгээний давалгаа хоорондын харьцуулалт")
    add_figure_source(doc, "Эх сурвалж: Оюутны тооцоолол")

    add_heading(doc, "4.4. Тогтвортой байдлын шинжилгээ", level=2)
    add_paragraph(doc,
        "Үндсэн үр дүнгийн тогтвортой байдлыг түүврийн дэд бүлгүүдээр "
        "шалгасан үр дүнг Хүснэгт 6-т нэгтгэн танилцуулав. Эмэгтэй "
        "ажилтнуудын боловсролын өгөөж (β = 0.0774) нь эрэгтэй "
        "ажилтнуудынхаас (β = 0.0565) 36%-иар өндөр гарсан нь Монгол "
        "Улсын хөдөлмөрийн зах зээлд хүйсээр ялгавартай бүтэц бий "
        "болохыг тод харуулж байна.")
    add_paragraph(doc,
        "Энэхүү хүйсийн ялгаа нь хоёр эдийн засгийн механизмаар "
        "тайлбарлагдах боломжтой. Нэгдүгээрт, Монгол Улсын хөдөлмөрийн "
        "зах зээлд мэргэжлийн хүйсээр ялгаатай тархалт бий: "
        "эмэгтэйчүүд нь боловсрол шаарддаг салбарт (эмнэлэг, "
        "сургууль, санхүү, үйлчилгээ) илүү төвлөрч ажилладаг тул "
        "нэмэлт боловсрол нь цалинд шууд илэрнэ. Хоёрдугаарт, бага "
        "боловсролтой эмэгтэйчүүд нь гэр ахуйн эдийн засаг, албан бус "
        "хөдөлмөрийн зах зээлд илүү оролцдог тул цалинтай ажилтны "
        "түүвэрт хамрагдахгүй байх сонголтын нөлөө үүсч болзошгүй. "
        "Ай Ар Ай Эм ХХК (2015)-ийн тайлан мөн ижил хандлагыг "
        "баталсан. Харин насны дэд бүлгүүдэд үнэлгээ ойролцоогоор "
        "0.067 түвшинд тогтвортой байгаа нь боловсролын өгөөж үе "
        "хоорондын хувьд системтэй ялгаагүй болохыг харуулна.")

    # Table 3.4
    add_table_caption(doc, "6", "Үндсэн үр дүнгийн робастнес шинжилгээ (OLS+FE)")
    build_table(doc,
        headers=["Дэд түүвэр", "β (боловсролын жил)", "Стандарт алдаа", "N"],
        rows=[
            ("Бүх түүвэр", "0.0661", "(0.0019)", "49,357"),
            ("Эрэгтэй", "0.0565", "(0.0026)", "24,788"),
            ("Эмэгтэй", "0.0774", "(0.0033)", "24,569"),
            ("Хот (Улаанбаатар)", "0.0649", "(0.0016)", "33,531"),
            ("Насны бүлэг 25–40", "0.0675", "(0.0020)", "27,879"),
            ("Насны бүлэг 41–60", "0.0681", "(0.0031)", "21,478"),
        ],
        col_widths_cm=[5.0, 3.7, 3.7, 3.0], align='right')
    add_table_source(doc, "Эх сурвалж: Оюутны тооцоолол")

    add_paragraph(doc,
        "Хэрэгсэл хувьсагчийн хүчин төгөлдөр байдлыг хяналтын тест "
        "(placebo)-аар нэмэлтээр шалгав. Боловсролын 8 жилээс доош "
        "(буюу ЕБС-ийн дунд ангиас доош) боловсролтой ажилтнуудын "
        "дэд түүвэрт Улаанбаатар хүртэлх зайн эхний шатны "
        "коэффициент нь t = 0.80 буюу ач холбогдолгүй гарсан. "
        "Харин 8 жилээс дээш боловсролтой дэд түүвэрт коэффициент "
        "t = 3.32 буюу ач холбогдолтой гарав. Энэ нь хэрэгсэл "
        "хувьсагч нь зөвхөн боловсролыг сонгох шийдвэрт нөлөөлдөг "
        "боловч бага боловсролтой иргэдийн цалинд шууд нөлөөлдөггүй "
        "гэсэн экзоген нөхцлийг эмпирикээр дэмжих болно.")
    add_paragraph(doc,
        "Түүврийн сонголтын хазайлтыг Хекманы хоёр шатлалт аргаар "
        "шалгасан. Милсийн урвуу харьцааны коэффициент 0.0406 "
        "(стандарт алдаа 0.0783, t = −0.52) буюу статистикийн хувьд "
        "ач холбогдолгүй байсан. Хекманы засвартай β_educ = 0.0671 "
        "нь энгийн OLS-ийн 0.0661-тай бараг адил гарсан. Үүнээс "
        "үзвэл цалинтай ажилтны түүврийн сонголт нь үндсэн үр "
        "дүнд мэдэгдэхүйц хазайлт үүсгээгүй болно.")

    add_heading(doc, "4.5. Эконометрик загваруудын нэгдсэн дүгнэлт", level=2)
    # Table 3.5
    add_table_caption(doc, "7", "Эконометрик загваруудын нэгдсэн дүгнэлт")
    build_table(doc,
        headers=["Загвар", "β (боловсрол)", "Жилийн өгөөж (%)", "N"],
        rows=[
            ("OLS энгийн", "0.0736", "7.6", "49,357"),
            ("OLS + аймаг, давалгаа FE", "0.0661", "6.8", "49,357"),
            ("Псевдо-панел FE", "0.1627", "17.7", "692"),
            ("Псевдо-панел RE", "0.2028", "22.5", "692"),
            ("2SLS IV (төрсөн аймаг)", "0.1133", "12.0", "12,020"),
            ("2SLS IV (ln d_UB)", "0.2055", "22.8", "7,746"),
            ("2SLS IV (overidentified)", "0.1043", "11.0", "7,746"),
            ("IVTR регим 1 (educ ≤ 13)", "0.0533", "5.5", "6,144"),
            ("IVTR регим 2 (educ > 13)", "0.1650", "17.9", "5,876"),
            ("Хекманы засвартай", "0.0671", "6.9", "49,811"),
        ],
        col_widths_cm=[6.5, 3.2, 3.5, 2.8], align='right')
    add_table_source(doc, "Эх сурвалж: Оюутны тооцоолол")

    add_paragraph(doc,
        "Хүснэгт 7-оос харвал OLS, псевдо-панел FE/RE, 2SLS, IVTR "
        "гэсэн дөрвөн шатлалт шинжилгээ нь нэгдмэл дүр зураг өгч "
        "байна. Хамгийн нарийвчилсан үнэлгээ болох IVTR нь "
        "боловсролын өгөөж шугаман биш, 13 жилийн босгоны хоёр талд "
        "ялгаатай бүтэцтэй болохыг баталсан. Энгийн OLS-ийн 6.8%-ийн "
        "дундаж нь уг бүтцийн хугарлыг далдалж буй нэгдсэн тоо байсан "
        "бөгөөд IVTR-ын регим 1 (5.5%) ба регим 2 (17.9%)-ийн дунд "
        "оршиж байгаа нь эдийн засгийн утгын хувьд цэгцтэй дүр зургийг "
        "гаргаж байна.")

    # ─── ТАВДУГААР БҮЛЭГ. ХЭЛЭЛЦҮҮЛЭГ ───────────────────────────────────────
    add_heading(doc, "V бүлэг. Хэлэлцүүлэг", level=1, before_break=True)

    add_heading(doc, "5.1. Гол олдворуудын эдийн засгийн тайлбар", level=2)
    add_paragraph(doc,
        "Энэхүү судалгааны гол эмпирик олдвор нь Монгол Улсын "
        "хөдөлмөрийн зах зээлд боловсролын өгөөж 13 жилийн босгон дээр "
        "эрс огцом үсрэлттэй байна гэдэг явдал юм. Уг босго нь санамсаргүй "
        "тохиолдсон зүйл биш, харин Монгол Улсын боловсролын "
        "тогтолцооны институцийн бүтцэд нягт холбоотой. 12 жилийн "
        "ерөнхий боловсрол нь 2008 оны шинэчлэлээс хойш нэгдсэн "
        "стандарт болж, 13-ны жил нь МСҮТ-ийн эхний жил, коллежийн "
        "нэг жилийн бэлтгэл сургалт, их сургуулийн анхны курс зэрэгтэй "
        "тохирч байна.")
    add_paragraph(doc,
        "Уг бүтцийн хугарал нь Jaeger ба Page (1996)-ийн \"дипломын "
        "нэмэгдэл\" үзэгдэл (sheepskin effect) болон Spence (1973)-ийн "
        "дохионы онолын (signaling) таамаглалтай бүрэн нийцэж байна. "
        "Ажил олгогчид боловсролыг жил тус бүрд нь биш, харин дипломын "
        "түвшнээр ялгаж, ажиллах хүчний чанарын дохио болгон ашигладаг. "
        "Монгол Улсад төрийн албанд орох шалгалт, эрүүл мэнд, "
        "боловсрол, хууль эрх зүй, санхүүгийн салбарт бакалаврын "
        "диплом хатуу шалгуур болдог тул диплом эзэмшсэн хүмүүс "
        "тусдаа өндөр цалингийн ангилалд хамаардаг. Мөн Ай Ар Ай Эм "
        "ХХК (2015)-ийн тайланд тэмдэглэснээр дээд боловсролтой "
        "иргэдийн ажилгүйдлийн түвшин (5.6%) бусад түвшнийхээс "
        "доогуур байдаг нь дипломын институцийн үнэ цэнийг "
        "баталгаажуулж байна.")
    add_paragraph(doc,
        "Мөн чухал нь, доод регимд илэрсэн 5.5%-ийн харьцангуй бага "
        "өгөөж нь Монголын ЕБС-ийн дунд шатны мэргэжлийн ур чадварын "
        "түвшинтэй уялдаатай юм. Ажиллах хүчний зах зээл дэх эрэлт "
        "хурдтай өөрчлөгдөж буй орчинд (мэдээллийн технологи, санхүү, "
        "гадаад хэл, нарийн мэргэжлийн техникийн ур чадвар), ЕБС-ийн "
        "сургалтын агуулга одоогийн ажил олгогчдын шаардлагад бүрэн "
        "нийцэж чадахгүй байна. Энэ нь Монголбанк-ны судлаач Болдбаатар "
        "(2017)-ын дурдсан боловсролын чанарыг бус жилийг хэмжих "
        "стандарт хандлагын сул талыг мөн илтгэж буй юм.")

    add_heading(doc, "5.2. Өмнөх судалгаануудтай харьцуулалт", level=2)
    add_paragraph(doc,
        "Энэхүү судалгааны OLS+FE үнэлгээ (6.8%) нь Pastore (2010)-ийн "
        "7–9%, Дэлхийн банкны Claudio ба Patrinos (2014)-ийн 10.1% "
        "тооцоотой ерөнхийдөө нийцэж байна. Гэхдээ уг өмнөх бүтээлүүд "
        "нь нэг нэгдсэн дундаж үнэлгээг гаргаж, босгоны хоёр талын "
        "ялгааг далдлан харуулаагүй. Иймд бодлогын тайланд \"Монголд "
        "боловсролын өгөөж дундаж 8% орчим\" гэсэн нэг тооны дүр "
        "зургийг бий болгосон нь бүх түвшний боловсролд ижил үр "
        "нөлөөтэй гэсэн төөрөгдүүлсэн ойлголтыг бий болгов.")
    add_paragraph(doc,
        "Энэхүү судалгааны IVTR үр дүнг Pastore-ийн жигнэсэн дундаж "
        "тооцооллоор харахад: хэрэв түүврийн 40% нь 13 жилээс доош, "
        "60% нь 13 жилээс дээш боловсролтой байсан гэж үзвэл, "
        "жигнэсэн дундаж нь 0.4 × 5.5% + 0.6 × 17.9% ≈ 12.9% байх "
        "ёстой. Энэ тоо OLS-д хязгаарлагдсан Pastore (2010)-ийн 7–9% "
        "үнэлгээнээс 40%-иар өндөр байх байсан. Зөрүү нь OLS-ийн "
        "доош чиглэсэн хэмжилтийн хазайлт болон босгыг огт авч "
        "үзээгүйгээс үүдсэн болох нь харагдаж байна.")
    add_paragraph(doc,
        "Олон улсын контекстэд Psacharopoulos ба Patrinos (2018)-ийн "
        "139 орныг хамарсан мета-шинжилгээнд боловсролын өгөөжийн "
        "дэлхийн дундаж жилд 9.0%, дээд боловсролынх 14.6% гэж "
        "тогтоосон. Манай IVTR-ын дээд регимийн 17.9% нь энэхүү "
        "олон улсын жишгээс даваад байна. Энэ нь хөгжиж буй орнуудын "
        "хөдөлмөрийн зах зээлийн онцлогт нийцдэг дүр зураг юм: "
        "дипломын хомсдлын үед итгэмжлэлийн үнэ цэнэ харьцангуй "
        "өндөр байдаг (Duflo, 2001).")

    add_heading(doc, "5.3. Судалгааны хязгаарлалтууд", level=2)
    add_paragraph(doc,
        "Энэхүү судалгаа хэд хэдэн хязгаарлалттай болохыг ил тод "
        "хүлээн зөвшөөрөх шаардлагатай. Нэгдүгээрт, шинжилгээ нь "
        "зөвхөн цалинтай ажилтнуудын түүверт үндэслэгдсэн тул өөрийн "
        "бизнес эрхлэгчид болон албан бус секторын ажилчдыг "
        "хамраагүй. Энэ нь β-г бага зэрэг дээш чиглэсэн хазайлтад "
        "орох магадлалтай гэж үзэж болно — учир нь өндөр боловсролтой "
        "хэрнээ өөрийн бизнесээр доогуур цалинтай байгаа ажилтнууд "
        "түүвэрт хамрагдаагүй. Хекманы засвар нь энэ хазайлтыг "
        "статистикийн хувьд ач холбогдолгүй гэж баталсан боловч "
        "үүнийг бүрэн арилгаж чадаагүй.")
    add_paragraph(doc,
        "Хоёрдугаарт, 2016 ба 2018 оны давалгаанд төрсөн аймгийн "
        "мэдээлэл байхгүй байсан тул үндсэн IV шинжилгээнд оруулж "
        "чадаагүй. Энэ нь түүврийн хэмжээг багасгаж стандарт алдааг "
        "өсгөж болох боловч коэффициентийн хазайлт үүсгэхгүй гэдгийг "
        "тогтвортой байдлын шинжилгээнд баталсан. Гуравдугаарт, "
        "боловсролын чанарын хувьсагч (шалгалтын оноо, сургуулийн "
        "нэр, мэргэжлийн ангилал) ӨНЭЗС-д байхгүй тул зөвхөн "
        "боловсролын жилийн тоог ашиглахаас өөр аргагүй байсан. "
        "Ур чадварын далд хазайлтыг хэрэгсэл хувьсагч хэсэгчлэн "
        "засч байгаа боловч чанарын дотоод хэмжлийн зөрөөг бүрэн "
        "барих боломжгүй. Дараагийн судалгаанд ӨНЭЗС-д боловсролын "
        "чанарын үзүүлэлт нэмэх нь эдгээр хязгаарлалтыг арилгах "
        "чухал алхам болно.")

    # ─── ЗУРГААДУГААР БҮЛЭГ. ДҮГНЭЛТ ─────────────────────────────────────────
    add_heading(doc, "VI бүлэг. Дүгнэлт ба бодлогын санал",
                level=1, before_break=True)

    add_heading(doc, "6.1. Үндсэн дүгнэлт", level=2)
    add_paragraph(doc,
        "Энэхүү судалгаа нь Монгол Улс дахь боловсролын бодит өгөөжийг "
        "Caner–Hansen (2004)-ийн хэрэгсэл хувьсагчтай босготой "
        "регрессийн аргаар үнэлэн, гурван гол дүгнэлтэд хүрэв.")
    add_paragraph(doc,
        "Нэгдүгээрт, Монгол Улс дахь боловсролын өгөөж шугаман бус, "
        "бүтцийн хугаралтай болохыг эмпирикээр тогтоов. OLS-ийн 6.8%-"
        "ийн дундаж үнэлгээ нь 13 жилийн босгоны хоёр талд орших "
        "ялгаатай регимүүдийн дундаж үр дүн байсан. Босгоноос доош "
        "5.5%, босгоноос дээш 17.9% буюу 3.3 дахин их ялгаа нь "
        "Монголын хөдөлмөрийн зах зээлд дипломын нэмэгдэл үзэгдэл "
        "оршин буйг эмпирикээр нотолж байна. Энэхүү олдвор нь Монгол "
        "Улс дахь анхны IVTR шинжилгээний үр дүн бөгөөд Jaeger ба "
        "Page (1996)-ийн \"дипломын нэмэгдэл\" онолыг Монгол нутгийн "
        "хөдөлмөрийн зах зээлд баталсан ажил болов.")
    add_paragraph(doc,
        "Хоёрдугаарт, эндогений хазайлтыг засах нь үр дүнтэй бөгөөд "
        "мэдэгдэхүйц өөрчлөлтөд хүргэж байна. OLS-ийн 6.8%-аас 2SLS-"
        "ийн 11.3% хүртэлх өсөлт буюу 49%-ийн ялгаа нь хэмжилтийн "
        "алдаанаас үүдэлтэй доош чиглэсэн хазайлт нь ур чадварын "
        "хомсдлоос үүдэлтэй дээш чиглэсэн хазайлтаас давамгайлж "
        "байгааг харуулж байна. Энэхүү дүгнэлт нь Card (1993)-ын "
        "Хойд Америкт хийсэн сонгодог үр дүнгийн Монгол Улсад "
        "хийгдсэн анхны баталгаажуулалт юм.")
    add_paragraph(doc,
        "Гуравдугаарт, уг эмпирик олдвор нь боловсролын бодлогын "
        "хувьд ихээхэн ач холбогдолтой. Бүрэн дунд боловсролыг "
        "дуусгаснаас дээд боловсролд шилжих шилжилтийн цэгт "
        "тусгайлан хөрөнгө оруулалт хийх нь Монгол Улсын ажиллах "
        "хүчний чадамжийг дээшлүүлэхэд хамгийн өндөр эргэн "
        "төлөлттэй бодлогын арга хэрэгсэл болох магадлалтай. "
        "Түүнчлэн Даваажаргал ба Цолмон (2019)-ийн макро түвшний "
        "Threshold SVAR шинжилгээтэй хамт Монгол Улсад босго "
        "шинжилгээний эмпирик уламжлалыг бий болгож байгаа нь "
        "дараагийн судлаачдад шинэ боломжийг нээх юм.")

    add_heading(doc, "6.2. Бодлогын санал", level=2)

    add_paragraph(doc,
        "Судалгааны эмпирик үр дүнд тулгуурлан Монгол Улсын "
        "боловсролын бодлого болон хөдөлмөрийн зах зээлийн "
        "зохицуулалтад хэрэгжүүлэх боломжтой гурван саналыг "
        "боловсруулав. Эдгээр нь IVTR-ын тоон үр дүнд суурилсан "
        "жишиг тооцоо бөгөөд цаашдын илүү нарийвчилсан тооцоог "
        "шаардана.")
    policy_items = [
        ("Дээд боловсролд хүрэх санхүүгийн саадыг бууруулах. ",
         "Босгоноос дээш нэг жилийн өгөөж 17.9% (доод регимээс 12.4 "
         "нэгж хувиар өндөр) гэсэн эмпирик тоон нотолгоонд тулгуурлан "
         "Засгийн газар оюутны зээл ба тэтгэлгийн нийлүүлэлтийг "
         "тэлэх нь хамгийн өндөр эргэн төлөлттэй бодлого болохыг "
         "нотолж байна. Тухайлбал, 3 жилийн бакалаврын сургалт нь "
         "бодит ажил ахуйн үнэ цэнээр ойролцоогоор 37%-ийн цалингийн "
         "өсөлтийг үүсгэж болох жишиг тооцоо гарч байна. Хөдөөгийн "
         "амжилттай төгсөгчид дипломт боловсролд хамрагдах саадыг "
         "бууруулах хөтөлбөрийг хэрэгжүүлэх шаардлагатай."),

        ("Сумын ЕБС-ийн чанарын ялгааг бууруулах. ",
         "Сумын ЕБС-ийн багш/сурагчийн харьцаа нь боловсролын "
         "гарцад хамгийн хүчтэй нөлөөтэй хэрэгсэл хувьсагч болсон "
         "(F = 110.02). Тиймээс уг харьцаа хангалтгүй буюу нэг "
         "багшид 30-аас дээш сурагч ногдож буй сумуудад багшийн "
         "цалин, мэргэшүүлэх сургалт, дэд бүтцийн тусгай хөтөлбөр "
         "хэрэгжүүлэх нь хүмүүн капиталын хуримтлалд өндөр "
         "өгөөжтэй байх боломжтой. Сум бүрт 2-3 багш нэмэх нь "
         "0.2-0.5 жилийн нэмэлт сургалтын хуримтлалыг үүсгэж 6-12%-"
         "ийн цалингийн өсөлтөд хүргэх магадлалтай (жишиг тооцоо)."),

        ("ӨНЭЗС-ийн боловсролын мониторингийг шинэчлэх. ",
         "ҮСХ-ны ӨНЭЗС-д дараах гурван шинэ хувьсагчийг нэмж оруулах "
         "нь дараагийн үеийн судалгааны боломжийг нээнэ: (а) ЭЕШ-ын "
         "оноо эсвэл боловсролын чанарын индикатор, (б) диплом "
         "олгосон сургуулийн нэр, (в) мэргэжлийн нарийвчилсан "
         "ангилал. Эдгээр мэдээллийн нэмэлт нь боловсролын "
         "чанарын нөлөөг зөвхөн жилийн тоогоор хэмжихийн оронд "
         "шууд үнэлэх, ЕБС ба их дээд сургуулиудыг чанараар "
         "харьцуулах эмпирик бааз суурийг бүрдүүлнэ."),
    ]
    for i, (title, body) in enumerate(policy_items, 1):
        add_mixed(doc, [
            (f"{i}) ", {'size': 12, 'bold': True}),
            (title, {'size': 12, 'bold': True}),
            (body, {'size': 12}),
        ], indent_first=True)

    add_paragraph(doc,
        "Дүгнэвэл, Монгол Улсын боловсрол нь хөдөлмөрийн орлогод эрс "
        "эерэг нөлөөтэй ч тэрхүү нөлөө нь боловсролын түвшинд жигд "
        "бус, 13 жилийн босгон дээр огцом үсрэлттэй. Ерөнхий "
        "боловсролоос дипломт боловсрол руу шилжих замд оршиж буй "
        "саадыг бууруулах, чанарын зөрүүг арилгах нь Монгол Улсын "
        "хүмүүн капиталын хуримтлалын хамгийн өндөр эргэн төлөлттэй "
        "хөрөнгө оруулалт байж болох юм. Цаашдын судалгаанд "
        "боловсролын чанарын хэмжлийг нэмж, олон улсын харьцуулалтыг "
        "өргөжүүлэх нь энэхүү ажлын үргэлжлэл байж болно.")

    # ─── НОМ ЗҮЙ ─────────────────────────────────────────────────────────────
    add_heading(doc, "Ном зүй", level=1, before_break=True)
    refs = [
        "Becker, G. S. (1964). *Human capital: A theoretical and empirical "
        "analysis with special reference to education*. National Bureau of "
        "Economic Research.",

        "Blundell, R., Dearden, L., & Sianesi, B. (2005). Evaluating the "
        "effect of education on earnings: Models, methods and results from "
        "the National Child Development Survey. *Journal of the Royal "
        "Statistical Society: Series A*, 168(3), 473–512.",

        "Caner, M., & Hansen, B. E. (2004). Instrumental variable "
        "estimation of a threshold model. *Econometric Theory*, 20(5), "
        "813–843.",

        "Card, D. (1993). *Using geographic variation in college proximity "
        "to estimate the return to schooling* (NBER Working Paper No. "
        "4483). National Bureau of Economic Research.",

        "Card, D. (1999). The causal effect of education on earnings. In "
        "O. Ashenfelter & D. Card (Eds.), *Handbook of Labor Economics* "
        "(Vol. 3A, pp. 1801–1863). Elsevier.",

        "Chakroun, M. (2013). Threshold effects in the relationship "
        "between financial development and income inequality. *Economics "
        "Letters*, 118(2), 378–384.",

        "Conley, T. G., Hansen, C. B., & Rossi, P. E. (2012). Plausibly "
        "exogenous. *Review of Economics and Statistics*, 94(1), 260–272.",

        "Dean, J. T., & Jayachandran, S. (2020). Changing family attitudes "
        "to promote female employment. *AEA Papers and Proceedings*, 110, "
        "201–205.",

        "Deaton, A. (1985). Panel data from time series of cross-sections. "
        "*Journal of Econometrics*, 30(1–2), 109–126.",

        "Duflo, E. (2001). Schooling and labor market consequences of "
        "school construction in Indonesia: Evidence from an unusual "
        "policy experiment. *American Economic Review*, 91(4), 795–813.",

        "Hansen, B. E. (2000). Sample splitting and threshold estimation. "
        "*Econometrica*, 68(3), 575–603.",

        "Hansen, H., & Tarp, F. (2001). Aid and growth regressions. "
        "*Journal of Development Economics*, 64(2), 547–570.",

        "Heckman, J. J. (1979). Sample selection bias as a specification "
        "error. *Econometrica*, 47(1), 153–161.",

        "IMF. (2023). *Mongolia: 2023 Article IV consultation* (IMF "
        "Country Report No. 23/348). International Monetary Fund.",

        "Jaeger, D. A., & Page, M. E. (1996). Degrees matter: New "
        "evidence on sheepskin effects in the returns to education. "
        "*Review of Economics and Statistics*, 78(4), 733–740.",

        "Jensen, R. (2010). The (perceived) returns to education and the "
        "demand for schooling. *Quarterly Journal of Economics*, 125(2), "
        "515–548.",

        "Khan, M. S., & Senhadji, A. S. (2001). Threshold effects in the "
        "relationship between inflation and growth. *IMF Staff Papers*, "
        "48(1), 1–21.",

        "Khandker, S. R. (2005). Microfinance and poverty: Evidence using "
        "panel data from Bangladesh. *World Bank Economic Review*, 19(2), "
        "263–286.",

        "Mincer, J. (1974). *Schooling, experience, and earnings*. "
        "National Bureau of Economic Research.",

        "Pastore, F. (2010). Returns to education of young people in "
        "Mongolia. *Post-Communist Economies*, 22(2), 247–265.",

        "Psacharopoulos, G., & Patrinos, H. A. (2018). Returns to "
        "investment in education: A decennial review of the global "
        "literature. *Education Economics*, 26(5), 445–458.",

        "Schultz, T. W. (1961). Investment in human capital. *American "
        "Economic Review*, 51(1), 1–17.",

        "Spence, M. (1973). Job market signaling. *Quarterly Journal of "
        "Economics*, 87(3), 355–374.",

        "Staiger, D., & Stock, J. H. (1997). Instrumental variables "
        "regression with weak instruments. *Econometrica*, 65(3), 557–586.",

        "Stock, J. H., & Yogo, M. (2005). Testing for weak instruments in "
        "linear IV regression. In D. W. K. Andrews & J. H. Stock (Eds.), "
        "*Identification and inference for econometric models* (pp. 80–"
        "108). Cambridge University Press.",

        "World Bank. (2013). *Mongolia: Poverty assessment*. World Bank.",

        "World Bank. (2024). *Mongolia economic update*. World Bank.",

        "Үндэсний статистикийн хороо. (2024). *Өрхийн нийгэм эдийн "
        "засгийн судалгаа 2024*. УСХ. https://data.1212.mn",
    ]
    for r in refs:
        parts = []
        # Split on italic markers
        tokens = r.split('*')
        for i, tok in enumerate(tokens):
            if tok:
                parts.append((tok, {'italic': i % 2 == 1, 'size': 11}))
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=3, after=6, line_spacing=1.15)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # Hanging indent
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.first_line_indent = Cm(-1.25)
        for text, props in parts:
            run = p.add_run(text)
            set_run_props(run, size=props['size'], italic=props.get('italic', False))

    # ─── ХАВСРАЛТ ────────────────────────────────────────────────────────────
    add_heading(doc, "Хавсралт", level=1, before_break=True)

    add_heading(doc, "Хавсралт А. Альтернатив босго утгууд", level=2)
    add_paragraph(doc,
        "Оновчтой босгыг тогтоохдоо Caner–Hansen алгоритмын сүлжээн "
        "хайлтад 12, 13, 14 гэсэн утгуудыг турших бөгөөд алдааны "
        "квадратуудын нийлбэрийг хамгийн бага болгодог утгыг γ* "
        "хэмээн сонгоно. Дараах хүснэгтэд альтернатив утгуудын үр "
        "дүнг харьцуулан танилцуулав.")

    add_table_caption(doc, "А.1", "Альтернатив босго утгуудын үр дүн")
    build_table(doc,
        headers=["γ (жил)", "β₁ регим 1", "Өгөөж (%)/жил регим 1",
                 "β₂ регим 2", "Өгөөж (%)/жил регим 2", "Ялгаа"],
        rows=[
            ("12", "0.0762", "7.9", "0.1729", "18.9", "11.0"),
            ("13 (сонгогдсон)", "0.0533", "5.5", "0.1650", "17.9", "12.4"),
            ("14", "0.1149", "12.2", "0.2095", "23.3", "11.1"),
        ],
        col_widths_cm=[2.5, 2.3, 3.2, 2.3, 3.2, 2.0], align='right')
    add_table_source(doc, "Эх сурвалж: Оюутны тооцоолол")

    add_heading(doc, "Хавсралт Б. 1212.mn нээлттэй өгөгдлийн бүтэц", level=2)
    add_paragraph(doc,
        "Хэрэгсэл хувьсагч Z₃ (сумын ЕБС-ийн багш/сурагчийн харьцаа)-г "
        "бүтээхэд дараах гурван хүснэгтийг Үндэсний статистикийн "
        "хорооны PX-Web JSON API-аас татав:")
    lst = [
        "DT_NSO_2001_001V1 — Сумын ЕБС-ийн тоо 2000–2024.",
        "DT_NSO_2001_002V1 — Сумын ЕБС-ийн багшийн тоо 2000–2024.",
        "DT_NSO_2001_004V1 — Сумын ЕБС-ийн сурагчийн тоо 2000–2024.",
    ]
    for item in lst:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=3, after=3)
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        run = p.add_run(f"— {item}")
        set_run_props(run, size=12)
    add_paragraph(doc,
        "Нийт 572 сум-жилийн ажиглалтыг нэгтгэн ӨНЭЗС-ийн "
        "хариулагчдын 12 настай байх жил, төрсөн сумтай холбон "
        "merge хийв.")

    add_heading(doc, "Хавсралт В. Программ хангамжийн код", level=2)
    add_paragraph(doc,
        "Бүх шинжилгээг R 4.4.3 дээр хийв. Үндсэн багцууд: fixest 0.11.2 "
        "(FE, 2SLS, кластер SE, Kleibergen-Paap F); plm 2.6-4 (Панел "
        "FE/RE + Хаусманы тест); AER 1.2-12 (IV cross-check); "
        "data.table 1.15.4 (Өгөгдлийн хэрэгсэл); haven 2.5.4 (SPSS "
        "файлыг унших); ggplot2 3.5.1 (Зурагжуулалт). Python-оор "
        "хийсэн хөндлөнгийн шалгалтанд pandas 2.2, linearmodels 6.0, "
        "statsmodels 0.14 ашиглав. Caner–Hansen IVTR-ын гараар "
        "хэрэгжүүлэлт 05_ivtr_caner_hansen.R скриптэд 180 орчим мөрт "
        "хэмжээтэй бичигдсэн.")

    # Tell Word to auto-update TOC / PAGEREF fields when the doc opens
    set_update_fields_on_open(doc)

    # Save
    doc.save(str(FINAL))
    print(f"\nSaved: {FINAL}")
    print(f"Size: {FINAL.stat().st_size:,} bytes")


if __name__ == "__main__":
    main()
